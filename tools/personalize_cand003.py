from pathlib import Path
from docx import Document

BASE = Path(__file__).resolve().parents[1] / "boveda-entrevista-profesional" / "busqueda-empleo" / "candidaturas" / "CAND-2026-003-randstad-auxiliar-administrativo-prl"

REPLACEMENTS = {
    "Candidato/a": "Gustavo Vega",
    "CANDIDATO/A": "GUSTAVO VEGA",
    "Datos de contacto pendientes de autorización para esta candidatura": "gusvegabol@gmail.com | 669 549 933",
    "Datos de contacto: pendientes de autorización para esta candidatura": "gusvegabol@gmail.com | 669 549 933",
}

def replace_in_paragraph(paragraph):
    text = paragraph.text
    for old, new in REPLACEMENTS.items():
        text = text.replace(old, new)
    if text != paragraph.text:
        for run in paragraph.runs:
            run.text = ""
        if paragraph.runs:
            paragraph.runs[0].text = text
        else:
            paragraph.add_run(text)

def apply(path):
    doc = Document(path)
    paragraphs = list(doc.paragraphs)
    for section in doc.sections:
        paragraphs += list(section.header.paragraphs) + list(section.footer.paragraphs)
    for paragraph in paragraphs:
        replace_in_paragraph(paragraph)
    if path.name == "carta-presentacion.docx":
        replacements = {
            6: "Me dirijo a ustedes para presentar mi candidatura al puesto de Auxiliar Administrativo/a PRL. Mi trayectoria reúne experiencia en dirección de RR. HH., gestión administrativa y documental laboral, organización de información, revisión de datos y mejora de procesos, con uso habitual de Word y Excel en la elaboración de informes.",
            7: "En Herfrailes, dentro de mis responsabilidades de RR. HH., estructuré documentación laboral, reforcé la trazabilidad y la protección de datos y diseñé flujos documentales para incorporación, vacaciones y permisos desde información organizada. También he trabajado con bases de datos, clasificación de información y generación de documentación oficial dentro de equipos. Cuento además con un curso breve de prevención de riesgos laborales realizado para mi desempeño profesional, aunque no he trabajado con plataformas CAE.",
            8: "Considero que mi experiencia en documentación laboral, gestión de información, precisión, organización y herramientas ofimáticas puede aportar valor al soporte administrativo del equipo de PRL. La duración temporal del contrato es una condición que me gustaría conocer con más detalle.",
        }
        for index, text in replacements.items():
            paragraph = doc.paragraphs[index]
            for run in paragraph.runs:
                run.text = ""
            paragraph.runs[0].text = text
    doc.save(path)

if __name__ == "__main__":
    apply(BASE / "cv.docx")
    apply(BASE / "carta-presentacion.docx")
