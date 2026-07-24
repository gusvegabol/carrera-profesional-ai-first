"""Aplica las reglas tipográficas comunes de Job-up a un DOCX existente."""
from pathlib import Path
import sys
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def is_structural(text):
    text = text.strip()
    return (not text or text.isupper() or text.startswith(("Candidatura:", "Datos de contacto", "RANDSTAD", "A la atención", "Estimado", "Atentamente,")) or text.endswith(",") or text[:4].isdigit())


def set_run(run):
    if run.font.size is None or run.font.size.pt < 10:
        run.font.size = Pt(10)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "es-ES")
    lang.set(qn("w:eastAsia"), "es-ES")
    lang.set(qn("w:bidi"), "es-ES")
    run._element.get_or_add_rPr().append(lang)


def apply(path):
    doc = Document(path)
    doc.styles["Normal"].font.size = Pt(10)
    for section in doc.sections:
        for para in list(section.header.paragraphs) + list(section.footer.paragraphs):
            for run in para.runs:
                set_run(run)
    for para in doc.paragraphs:
        for run in para.runs:
            set_run(run)
        if not is_structural(para.text):
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    doc.save(path)


if __name__ == "__main__":
    for argument in sys.argv[1:]:
        apply(Path(argument))
