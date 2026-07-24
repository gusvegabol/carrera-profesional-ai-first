from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "boveda-entrevista-profesional" / "busqueda-empleo" / "candidaturas" / "CAND-2026-005-proapro-tecnico-distribucion-logistica"

def style(run, size=10, bold=False, color="1F2937"):
    run.font.name = "Arial"; run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(max(size, 10)); run.font.bold = bold; run.font.color.rgb = RGBColor.from_string(color)
    lang = OxmlElement("w:lang"); lang.set(qn("w:val"), "es-ES"); lang.set(qn("w:eastAsia"), "es-ES"); lang.set(qn("w:bidi"), "es-ES")
    run._element.get_or_add_rPr().append(lang)

def para(doc, text="", size=10, bold=False, color="1F2937", after=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after); p.paragraph_format.line_spacing = 1.08
    p.alignment = align
    if text: style(p.add_run(text), size, bold, color)
    return p

def rule(p):
    borders = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8"); bottom.set(qn("w:color"), "2E74B5")
    borders.append(bottom); p._p.get_or_add_pPr().append(borders)

def main():
    doc = Document(); s = doc.sections[0]
    s.top_margin = Inches(.7); s.bottom_margin = Inches(.65); s.left_margin = Inches(.8); s.right_margin = Inches(.8)
    p = para(doc, "Gustavo Vega", 20, True, after=0, align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, "Candidatura: Técnico/a de Distribución y Logística — Pro a Pro", 11, False, "2E74B5", after=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, "gusvegabol@gmail.com | 669 549 933", 10, False, "5B6573", after=18, align=WD_ALIGN_PARAGRAPH.LEFT)
    h = para(doc, "Carta de presentación", 11, True, after=2, align=WD_ALIGN_PARAGRAPH.LEFT); rule(h)
    para(doc, "Estimado equipo de selección de Pro a Pro:", after=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, "Me dirijo a ustedes para presentar mi candidatura al puesto de Técnico/a de Distribución y Logística. Mi experiencia combina control de stock, análisis de rotación, planificación de movimientos entre centros y mejora de procesos, con el objetivo de sostener una distribución fiable, eficiente y orientada al servicio.", after=10)
    para(doc, "En Herfrailes S. L. diseñé un algoritmo de pedido a partir de ventas diarias que redujo caducidades un 30 % y productos sin venta durante más de un mes un 80 %. Posteriormente programé un sistema de redistribución entre tres tiendas: analizaba stock y rotación, generaba las instrucciones diarias de traslado y optimizaba la ruta de vehículos propios para que la mercancía pudiera reponerse con agilidad.", after=10)
    para(doc, "También definí reglas de surtido, analicé desviaciones y creé un seguimiento del mantenimiento de vehículos. Esta experiencia me ha acostumbrado a conectar datos, operaciones y continuidad del servicio con rigor, trazabilidad y atención al detalle.", after=10)
    para(doc, "Me atrae especialmente el enfoque de Pro a Pro en un servicio integral, fiable y eficiente para la hostelería organizada, apoyado en la calidad, la mejora continua y la atención a las necesidades de cada cliente. Considero que mi experiencia en control de stock, trazabilidad y mejora operativa puede contribuir a esa forma de trabajar.", after=10)
    para(doc, "Quedo a su disposición para ampliar mi experiencia en una entrevista.", after=10)
    para(doc, "Atentamente,", after=2, align=WD_ALIGN_PARAGRAPH.LEFT); para(doc, "Gustavo Vega", 10, True, after=0, align=WD_ALIGN_PARAGRAPH.LEFT)
    OUT.mkdir(parents=True, exist_ok=True); doc.save(OUT / "carta-presentacion.docx")

if __name__ == "__main__": main()
