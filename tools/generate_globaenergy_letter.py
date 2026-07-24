from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "boveda-entrevista-profesional" / "busqueda-empleo" / "candidaturas" / "CAND-2026-004-globaenergy-auxiliar-administrativo-back-office"

def run_style(run, size=10, bold=False, color="1F2937"):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(max(size, 10))
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    lang = OxmlElement("w:lang"); lang.set(qn("w:val"), "es-ES"); lang.set(qn("w:eastAsia"), "es-ES"); lang.set(qn("w:bidi"), "es-ES")
    run._element.get_or_add_rPr().append(lang)

def para(doc, text="", size=10, bold=False, color="1F2937", after=8, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    p.alignment = align
    if text:
        run_style(p.add_run(text), size, bold, color)
    return p

def rule(p):
    ppr = p._p.get_or_add_pPr(); borders = OxmlElement("w:pBdr"); bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "8"); bottom.set(qn("w:color"), "2E74B5")
    borders.append(bottom); ppr.append(borders)

def main():
    doc = Document(); s = doc.sections[0]
    s.top_margin = Inches(.7); s.bottom_margin = Inches(.65); s.left_margin = Inches(.8); s.right_margin = Inches(.8)
    p = para(doc, "Gustavo Vega", 20, True, after=0, align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, "Candidatura: Auxiliar Administrativo/a Back Office — Globaenergy", 11, False, "2E74B5", after=3, align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, "gusvegabol@gmail.com | 669 549 933", 10, False, "5B6573", after=18, align=WD_ALIGN_PARAGRAPH.LEFT)
    h = para(doc, "Carta de presentación", 11, True, after=2, align=WD_ALIGN_PARAGRAPH.LEFT); rule(h)
    para(doc, "Estimado equipo de selección de Globaenergy:", after=10, align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, "Me dirijo a ustedes para presentar mi candidatura al puesto de Auxiliar Administrativo/a Back Office. Mi trayectoria combina gestión administrativa y documental, organización de información, seguimiento de tareas y mejora de procesos, competencias que puedo trasladar al apoyo riguroso de la gestión de expedientes y contribuir a un servicio cercano y profesional.", after=10)
    para(doc, "En Herfrailes S. L. estructuré documentación, validé información, definí accesos y automaticé flujos documentales, reduciendo tareas manuales y mejorando la trazabilidad. También implanté sistemas de registro y seguimiento para coordinar varias áreas. En etapas anteriores trabajé con bases de datos, documentación oficial, contabilidad y gestión comercial.", after=10)
    para(doc, "Estoy habituado a trabajar con Word y Excel con nivel alto, correo electrónico, bases de datos y herramientas de seguimiento. Mi forma de trabajar se apoya en el rigor, el orden, la atención al detalle y un seguimiento fiable de cada expediente hasta su cierre.", after=10)
    para(doc, "Me interesa incorporarme a un entorno estable, comprometido con la calidad y la mejora continua, en el que pueda aportar esta experiencia administrativa y seguir desarrollándome. Quedo a su disposición para ampliar cualquier información en una entrevista.", after=10)
    para(doc, "Atentamente,", after=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, "Gustavo Vega", 10, True, after=0, align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.save(OUT / "carta-presentacion.docx")

if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True); main()
