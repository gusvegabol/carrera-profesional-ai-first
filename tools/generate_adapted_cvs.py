from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "boveda-entrevista-profesional" / "busqueda-empleo" / "candidaturas"
BLUE = "2E74B5"
INK = "1F2937"
MUTED = "5B6573"
HEADER_ALIGNMENT = WD_ALIGN_PARAGRAPH.CENTER
HEADER_SPACER_BEFORE_PT = 6.0
HEADER_SPACER_AFTER_PT = 2.0


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_bottom_border(paragraph, color=BLUE, size="8"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def set_run(run, size=10.0, bold=False, color=INK, italic=False):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(max(size, 10.0))
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "es-ES")
    lang.set(qn("w:eastAsia"), "es-ES")
    lang.set(qn("w:bidi"), "es-ES")
    run._element.get_or_add_rPr().append(lang)


def paragraph(doc, text="", size=10.0, bold=False, color=INK, italic=False, after=2, before=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = 1.0
    if align is not None:
        p.alignment = align
    if text:
        set_run(p.add_run(text), size, bold, color, italic)
    return p


def heading(doc, text):
    p = paragraph(doc, text.upper(), size=10.2, bold=True, color=INK, after=2, before=6, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_bottom_border(p)
    return p


def bullet(doc, text, size=10.0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.16)
    p.paragraph_format.first_line_indent = Inches(-0.10)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    set_run(p.add_run(text), size)
    return p


def role(doc, title, meta, bullets):
    p = paragraph(doc, title, size=10.0, bold=True, color=INK, after=0, before=2, align=WD_ALIGN_PARAGRAPH.LEFT)
    p = paragraph(doc, meta, size=10.0, color=MUTED, italic=True, after=1, align=WD_ALIGN_PARAGRAPH.LEFT)
    for text in bullets:
        bullet(doc, text, 10.0)


def build(path, config):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.42)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.24)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.0

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.paragraph_format.space_before = Pt(0)
    set_run(footer.add_run(config["footer"]), 10.0, False, MUTED)

    paragraph(doc, config["name"], size=20, bold=True, color=INK, after=0, align=HEADER_ALIGNMENT)
    paragraph(doc, config["title"], size=10, color=INK, after=1, align=HEADER_ALIGNMENT)
    paragraph(doc, config["contact"], size=10, color=MUTED, after=5, align=HEADER_ALIGNMENT)
    paragraph(doc, "", before=HEADER_SPACER_BEFORE_PT, after=HEADER_SPACER_AFTER_PT, align=WD_ALIGN_PARAGRAPH.LEFT)

    heading(doc, "Perfil profesional")
    paragraph(doc, config["profile"], size=10, after=3)

    heading(doc, "Competencias clave")
    for item in config["competencies"]:
        bullet(doc, item, 10)

    heading(doc, "Experiencia profesional relevante")
    for item in config["roles"]:
        role(doc, item["title"], item["meta"], item["bullets"])

    heading(doc, "Aptitudes técnicas y administrativas")
    for label, value in config["aptitudes"]:
        p = paragraph(doc, after=1)
        set_run(p.add_run(label + ": "), 10, True, INK)
        set_run(p.add_run(value), 10, False, INK)

    heading(doc, "Formación")
    for item in config["training"]:
        bullet(doc, item, 10)

    heading(doc, config["value_heading"])
    for item in config["value"]:
        bullet(doc, item, 10)

    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


COMMON_ROLES = [
    {
        "title": "Director Ejecutivo — Herfrailes S. L.",
        "meta": "2011–2024 | Administración, organización, sistemas, operaciones y logística",
        "bullets": [
            "Organicé documentación e información, diferencié datos personales y profesionales y definí accesos según las necesidades administrativas.",
            "Automaticé flujos documentales desde información estructurada, reduciendo una preparación de unos 20 minutos a un proceso inmediato y trazable.",
            "Implanté sistemas de registro y seguimiento para centralizar información, coordinar áreas y controlar tareas.",
        ],
    },
    {
        "title": "Programador — INERZA, S. A.",
        "meta": "2009–2011 | Área de Vivienda del Gobierno de Canarias",
        "bullets": [
            "Diseñé bases de datos, clasifiqué información de baremación y generé documentación oficial asociada a cada solicitud dentro de un equipo.",
        ],
    },
    {
        "title": "Analista programador — Computerland Las Palmas / COMPAZ",
        "meta": "1985–1990 | Sistemas empresariales, contabilidad y gestión comercial",
        "bullets": [
            "Analicé y mejoré soluciones de contabilidad y gestión comercial según las necesidades del negocio.",
        ],
    },
]


def main():
    build(OUT / "CAND-2026-002-islas-natura-07-tecnico-administrativo" / "cv.docx", {
        "name": "Gustavo Vega",
        "title": "Técnico/a Administrativo/a | Gestión documental, datos y mejora de procesos",
        "contact": "gusvegabol@gmail.com | 669 549 933",
        "footer": "Gustavo Vega | CV Administrativo/a",
        "profile": "Trayectoria en gestión administrativa y documental, organización de información, bases de datos y mejora de procesos. Experiencia en Word, Excel, trazabilidad y automatización de tareas, con orientación al detalle y a los resultados.",
        "competencies": [
            "Gestión documental, archivo y validación de información.",
            "Organización de expedientes, prioridades y seguimiento.",
            "Word y Excel con nivel alto; Office 365 y Power BI previo.",
            "Bases de datos, análisis de información y control de procesos.",
            "Trazabilidad documental, protección de datos y control de accesos.",
            "Comunicación escrita, coordinación y mejora continua.",
        ],
        "roles": COMMON_ROLES,
        "aptitudes": [
            ("Ofimática y gestión", "Word, Excel, Office 365, gestión documental, Trello y Notion."),
            ("Procesos y administración", "Organización, seguimiento, contabilidad, gestión comercial y mejora de procesos."),
            ("Sistemas y datos", "SQL Server, bases de datos, Visual Studio, VB.NET y Windows Server."),
        ],
        "training": [
            "Máster en Inteligencia Artificial con certificación universitaria — BIG school; en curso.",
            "Formación Profesional en Administración de Empresas y Gestión Contable; no finalizada.",
            "Power BI — curso de Udemy y experiencia previa; uso no reciente.",
        ],
        "value_heading": "Valor diferencial para un puesto administrativo",
        "value": [
            "Capacidad para ordenar información, documentar procesos y convertir tareas manuales en flujos trazables.",
            "Experiencia combinando administración, datos y mejora continua con atención al detalle.",
        ],
    })

    build(OUT / "CAND-2026-003-randstad-auxiliar-administrativo-prl" / "cv.docx", {
        "name": "Candidato/a",
        "title": "Auxiliar Administrativo/a PRL | Gestión documental y ofimática",
        "contact": "Datos de contacto pendientes de autorización para esta candidatura",
        "footer": "CV Administrativo/a",
        "profile": "Experiencia en dirección de RR. HH., gestión administrativa y documental laboral, revisión y validación de información y organización de datos. Uso habitual de Word y Excel en informes, con especial atención a la trazabilidad, la precisión y el control de accesos.",
        "competencies": [
            "Gestión, archivo y validación de documentación.",
            "Revisión de información, expedientes y estados.",
            "Word y Excel con nivel alto; Office 365.",
            "Introducción, clasificación y control de datos.",
            "Trazabilidad documental y protección de datos.",
            "Documentación laboral, RR. HH. y control de accesos.",
            "Organización, trabajo en equipo y seguimiento.",
        ],
        "roles": [
            {
                "title": "Director Ejecutivo — Herfrailes S. L.",
                "meta": "2011–2024 | Dirección de RR. HH., administración, sistemas y organización",
                "bullets": [
                    "En funciones de RR. HH., estructuré documentación laboral, separé datos personales y profesionales y definí accesos según las necesidades administrativas.",
                    "Diseñé y automaticé flujos de documentación, incorporación, vacaciones y permisos desde información estructurada, con trazabilidad y control de datos.",
                    "Elaboré documentación e informes y mantuve sistemas de registro y seguimiento para coordinar tareas entre áreas.",
                ],
            },
            {
                "title": "Programador — INERZA, S. A.",
                "meta": "2009–2011 | Área de Vivienda del Gobierno de Canarias",
                "bullets": [
                    "Diseñé bases de datos, clasifiqué información de baremación y generé documentación oficial asociada a cada solicitud dentro de un equipo.",
                ],
            },
        ],
        "aptitudes": [
            ("Ofimática y gestión", "Word y Excel con nivel alto, Office 365, gestión documental, Trello y Notion."),
            ("RR. HH. y procesos", "Documentación laboral, organización, revisión, seguimiento, precisión y mejora de procesos."),
            ("Sistemas y datos", "SQL Server, bases de datos, Visual Studio, Windows 11 y Windows Server."),
        ],
        "training": [
            "Curso breve de prevención de riesgos laborales (PRL), realizado para el desempeño profesional; sin experiencia CAE.",
            "Formación Profesional en Administración de Empresas y Gestión Contable; no finalizada.",
            "Máster en Inteligencia Artificial con certificación universitaria — en curso.",
        ],
        "value_heading": "Valor diferencial para soporte administrativo PRL",
        "value": [
            "Experiencia transferible en revisión documental, control de información y trazabilidad.",
            "Curso breve de PRL, sin atribuir experiencia en plataformas de coordinación empresarial.",
        ],
    })

    build(OUT / "CAND-2026-004-globaenergy-auxiliar-administrativo-back-office" / "cv.docx", {
        "name": "Gustavo Vega",
        "title": "Auxiliar Administrativo/a Back Office | Gestión documental, expedientes y datos",
        "contact": "gusvegabol@gmail.com | 669 549 933",
        "footer": "Gustavo Vega | CV Back Office Administrativo/a",
        "profile": "Experiencia en gestión administrativa y documental, revisión de información, organización de expedientes y seguimiento de procesos. Manejo alto de Word y Excel, bases de datos y herramientas de trabajo colaborativo. Aporto rigor, orden y seguimiento fiable para sostener tareas administrativas simultáneas y contribuir a un servicio cercano y profesional en entornos de back office.",
        "competencies": [
            "Gestión documental, archivo y validación de información.",
            "Organización y seguimiento de expedientes y tareas.",
            "Word y Excel con nivel alto; Office 365 y correo electrónico.",
            "Bases de datos, clasificación de información y trazabilidad.",
            "Gestión comercial, facturación y apoyo a procesos administrativos.",
            "Comunicación escrita, coordinación y atención al detalle.",
        ],
        "roles": [
            {
                "title": "Director Ejecutivo — Herfrailes S. L.",
                "meta": "2011–2024 | Administración, documentación, sistemas y organización",
                "bullets": [
                    "Organicé y validé documentación administrativa desde información estructurada, separando datos personales y profesionales y definiendo accesos según la necesidad de cada área.",
                    "Diseñé y automaticé flujos documentales: una preparación manual de unos 20 minutos pasó a generarse de forma inmediata y trazable.",
                    "Implanté un circuito de registro y seguimiento para centralizar información, asignar tareas y facilitar el control de su estado entre áreas.",
                ],
            },
            {
                "title": "Programador — INERZA, S. A.",
                "meta": "2009–2011 | Área de Vivienda del Gobierno de Canarias",
                "bullets": [
                    "Clasifiqué información de expedientes, diseñé bases de datos y generé documentación oficial asociada a cada solicitud dentro de un equipo.",
                ],
            },
            {
                "title": "Analista programador — Computerland Las Palmas / COMPAZ",
                "meta": "1985–1990 | Contabilidad, gestión comercial y sistemas empresariales",
                "bullets": [
                    "Analicé necesidades de negocio y mejoré soluciones de contabilidad y gestión comercial.",
                ],
            },
        ],
        "aptitudes": [
            ("Ofimática y comunicación", "Word y Excel con nivel alto, Office 365, correo electrónico y elaboración de informes."),
            ("Administración y seguimiento", "Gestión documental, expedientes, archivo, facturación, gestión comercial y control de tareas."),
            ("Sistemas y datos", "SQL Server, bases de datos, Visual Studio, VB.NET, Trello y Notion."),
        ],
        "training": [
            "Formación Profesional en Administración de Empresas y Gestión Contable; no finalizada.",
            "Máster en Inteligencia Artificial con certificación universitaria — en curso.",
        ],
        "value_heading": "Valor diferencial para back office administrativo",
        "value": [
            "Capacidad para revisar documentación, ordenar expedientes y mantener el seguimiento de varias tareas con trazabilidad.",
            "Experiencia convirtiendo información estructurada en procesos administrativos más ágiles y controlables.",
        ],
    })

    build(OUT / "CAND-2026-005-proapro-tecnico-distribucion-logistica" / "cv.docx", {
        "name": "Gustavo Vega",
        "title": "Técnico/a de Distribución y Logística | Stock, rutas y mejora operativa",
        "contact": "gusvegabol@gmail.com | 669 549 933",
        "footer": "Gustavo Vega | CV Distribución y Logística",
        "profile": "Experiencia en operaciones de distribución alimentaria, análisis de stock, planificación de movimientos entre centros y mejora continua. Manejo alto de Excel, bases de datos y automatización de procesos. Aporto rigor operativo, trazabilidad y orientación a un servicio de distribución fiable y eficiente.",
        "competencies": [
            "Control de stock, análisis de rotación y acciones correctivas.",
            "Distribución entre centros, rutas y coordinación con vehículos propios.",
            "Planificación de pedidos, aprovisionamiento y reducción de caducidades.",
            "Excel con nivel alto, bases de datos e informes de actividad.",
            "Trazabilidad, análisis de desviaciones y mejora continua.",
            "Coordinación operativa, organización y atención al detalle.",
        ],
        "roles": [
            {
                "title": "Director Ejecutivo — Herfrailes S. L.",
                "meta": "2011–2024 | Operaciones de supermercados vinculados a SPAR en Gran Canaria",
                "bullets": [
                    "Diseñé una base de datos de ventas y un algoritmo de pedido que redujo un 30 % las caducidades y un 80 % los productos sin venta durante más de un mes.",
                    "Programé un sistema de redistribución de stock entre tres tiendas según rotación; generaba listas de preparación y optimizaba la ruta de vehículos propios para dos ciclos diarios de entrega.",
                    "Definí reglas de surtido y rotación, reduciendo stock inmovilizado y acortando el retorno de la inversión en compras.",
                    "Creé un sistema de seguimiento del mantenimiento de los vehículos de la empresa.",
                ],
            },
            {
                "title": "Analista programador — Granintra S. A.",
                "meta": "1990–1996 | Operaciones portuarias, consignataria, naviera y áridos",
                "bullets": [
                    "Implanté un sistema para operaciones de consignataria, naviera y carga de áridos, dentro del plazo de un año.",
                    "Desarrollé funciones de seguimiento de camiones que accedían y cargaban en el muelle.",
                ],
            },
        ],
        "aptitudes": [
            ("Operaciones y logística", "Stock, rotación, rutas, distribución, aprovisionamiento y mantenimiento de vehículos."),
            ("Datos y mejora", "Excel con nivel alto, bases de datos, análisis de desviaciones, automatización e informes."),
            ("Coordinación", "Organización operativa, seguimiento de tareas, trazabilidad y mejora continua."),
        ],
        "training": [
            "Formación Profesional en Administración de Empresas y Gestión Contable; no finalizada.",
            "Máster en Inteligencia Artificial con certificación universitaria — en curso.",
        ],
        "value_heading": "Valor diferencial para distribución y logística",
        "value": [
            "Experiencia conectando datos de ventas, stock y rotación con decisiones operativas de reposición y distribución.",
            "Capacidad para sostener una operación rigurosa, trazable y orientada a la continuidad del servicio.",
        ],
    })


if __name__ == "__main__":
    main()
