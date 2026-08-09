"""Generador determinista de documentos para candidaturas por oferta.

La IA aporta los textos finales en datos-generacion.json. Este módulo solo
valida, sustituye, convierte, verifica y publica los artefactos.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
from datetime import datetime, timezone
from copy import deepcopy
from pathlib import Path, PurePosixPath
from typing import Any, Iterable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image


SCRIPT_DIR = Path(__file__).resolve().parent
if str(SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(SCRIPT_DIR))

from componer_cv import construir_modelo_cv, renderizar_docx, renderizar_latex
TEMPLATE_ROOT = Path("boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas")
CANDIDATURE_ROOT = Path("boveda-entrevista-profesional/busqueda-empleo/candidaturas")
ERROR_ROOT = Path("boveda-entrevista-profesional/busqueda-empleo/registros-generacion")
TEMP_ROOT = Path(".tmp/job-up-generador")
DOCX_TO_PDF_TIMEOUT = 60

CV_FIELDS = (
    "[NOMBRE]", "[TITULAR]", "[EMAIL]", "[TELÉFONO]", "[LINKEDIN]",
    "[PERFIL PROFESIONAL]", "[PROPUESTA DE VALOR]",
    *[f"[EXPERIENCIA {i} CABECERA]" for i in range(1, 7)],
    *[f"[EXPERIENCIA {i} DESCRIPCION]" for i in range(1, 7)],
    *[f"[COMPETENCIA {i}]" for i in range(1, 5)],
    *[f"[FORMACION {i}]" for i in range(1, 4)],
    "[INFORMACION ADICIONAL]",
)

# Keep the order used by the template for deterministic JSON validation.
CV_FIELDS = (
    "[NOMBRE]", "[TITULAR]", "[EMAIL]", "[TELÉFONO]", "[LINKEDIN]",
    "[PERFIL PROFESIONAL]", "[PROPUESTA DE VALOR]",
    *sum(([f"[EXPERIENCIA {i} CABECERA]", f"[EXPERIENCIA {i} DESCRIPCION]"] for i in range(1, 7)), []),
    *[f"[COMPETENCIA {i}]" for i in range(1, 5)],
    *[f"[FORMACION {i}]" for i in range(1, 4)],
    "[INFORMACION ADICIONAL]",
)
CARTA_FIELDS = (
    "[NOMBRE]", "[TITULAR]", "[EMAIL]", "[TELÉFONO]", "[LINKEDIN]",
    "[DESTINATARIO]", "[FECHA]", "[ASUNTO]", "[SALUDO]", "[APERTURA]",
    "[EVIDENCIA 1]", "[EVIDENCIA 2]", "[ENCAJE]", "[CIERRE]",
    "[DESPEDIDA]", "[FIRMA]",
)
LATEX_FIELDS = CV_FIELDS
OUTPUT_KEYS = ("cv_docx", "cv_pdf", "cv_tex")
EXPECTED_TEMPLATES = {
    "template_cv": TEMPLATE_ROOT / "TEMPLATE_CV_FORMATO.docx",
    "template_carta": TEMPLATE_ROOT / "TEMPLATE_CARTA_PRESENTACION_FORMATO.docx",
    "template_latex": TEMPLATE_ROOT / "TEMPLATE_CV_FORMATO.tex",
}
EXPERIENCE_RE = re.compile(r"\[EXPERIENCIA ([1-6]) (CABECERA|DESCRIPCION)\]")
MARKER_RE = re.compile(r"\[[^\]]+\]")


def canonical_output_paths(route: Path) -> dict[str, Path]:
    return {
        "cv_docx": route / "cv.docx",
        "cv_pdf": route / "cv.pdf",
        "cv_tex": route / "cv.tex",
    }


def _project_relative(root: Path, value: str) -> Path:
    if not isinstance(value, str) or not value or "\\" in value:
        raise ValueError("Las rutas deben ser cadenas relativas con separadores '/'.")
    pure = PurePosixPath(value)
    if pure.is_absolute() or ".." in pure.parts or re.match(r"^[A-Za-z]:", value):
        raise ValueError(f"Ruta no permitida: {value}")
    unresolved = root / Path(*pure.parts)
    _reject_reparse_points(unresolved, value)
    result = unresolved.resolve(strict=False)
    root_resolved = root.resolve()
    try:
        result.relative_to(root_resolved)
    except ValueError as exc:
        raise ValueError(f"La ruta escapa de RUTA_PROYECTO: {value}") from exc
    return result


def _reject_reparse_points(path: Path, label: str) -> None:
    for candidate in (path, *path.parents):
        if not candidate.exists():
            continue
        if candidate.is_symlink() or getattr(candidate, "is_junction", lambda: False)():
            raise ValueError(f"La ruta contiene un enlace simbólico o junction no permitido: {label}")


def _assert_inside(path: Path, root: Path, label: str) -> None:
    try:
        path.resolve(strict=False).relative_to(root.resolve())
    except ValueError as exc:
        raise ValueError(f"{label} fuera de su raíz autorizada: {path}") from exc


def resolve_input_json(root: Path, value: str) -> Path:
    candidate = Path(value)
    if candidate.is_absolute():
        resolved = candidate.resolve(strict=False)
        _assert_inside(resolved, root, "datos-generacion.json")
        _reject_reparse_points(candidate, value)
        return resolved
    # El argumento del CLI admite la convención nativa de Windows. Las rutas
    # persistidas en JSON siguen validándose mediante _project_relative, que
    # exige '/' como separador canónico.
    return _project_relative(root, value.replace("\\", "/"))


def _assert_slot_map(section: Any, expected: tuple[str, ...], label: str) -> None:
    if not isinstance(section, dict) or tuple(section) != expected:
        raise ValueError(f"Contrato de marcadores inválido en {label}.")
    for key, value in section.items():
        if not isinstance(value, str) or any(ord(ch) < 32 for ch in value):
            raise ValueError(f"Valor inválido en {label}.{key}.")
        if "\r" in value or "\n" in value:
            raise ValueError(f"Los slots no admiten saltos de línea: {label}.{key}.")


def _validate_experiences(values: dict[str, str]) -> None:
    required = {"[NOMBRE]", "[TITULAR]", "[PERFIL PROFESIONAL]", "[PROPUESTA DE VALOR]"}
    for key in required:
        if not values[key].strip():
            raise ValueError(f"Campo obligatorio vacío: {key}")
    for i in range(1, 7):
        header = values[f"[EXPERIENCIA {i} CABECERA]"].strip()
        description = values[f"[EXPERIENCIA {i} DESCRIPCION]"].strip()
        if bool(header) != bool(description):
            raise ValueError(f"Pareja de experiencia incompleta: {i}")


def validate_payload(payload: dict[str, Any], project_root: Path) -> Path:
    if not isinstance(payload, dict) or payload.get("schema_version") != "1.0":
        raise ValueError("schema_version debe ser '1.0'.")
    expected_top = ("schema_version", "id_candidatura", "ruta_candidatura", "entradas", "salidas", "cv", "carta", "latex")
    if tuple(payload) != expected_top:
        raise ValueError("El JSON debe usar exactamente las claves del contrato 1.0.")
    candidate_id = payload["id_candidatura"]
    if not isinstance(candidate_id, str) or not re.fullmatch(r"CAND-\d{4}-\d{3}", candidate_id):
        raise ValueError("id_candidatura inválido.")
    route = _project_relative(project_root, payload["ruta_candidatura"])
    if not route.name.startswith(candidate_id):
        raise ValueError("La carpeta de candidatura no comienza por id_candidatura.")
    _assert_inside(route, (project_root / CANDIDATURE_ROOT), "ruta_candidatura")

    entries = payload["entradas"]
    if tuple(entries) != ("template_cv", "template_carta", "template_latex", "foto"):
        raise ValueError("Entradas inválidas.")
    for key, relative in EXPECTED_TEMPLATES.items():
        if entries[key] != relative.as_posix():
            raise ValueError(f"Plantilla no canónica: {key}")
        resolved = _project_relative(project_root, entries[key])
        if resolved != (project_root / relative).resolve():
            raise ValueError(f"Ruta de plantilla no canónica: {key}")
        if not resolved.is_file():
            raise ValueError(f"Falta la plantilla: {resolved}")
    photo = _project_relative(project_root, entries["foto"])
    if not photo.is_file():
        raise ValueError(f"Falta la fotografía: {photo}")
    _validate_photo(photo)

    if tuple(payload["salidas"]) != OUTPUT_KEYS:
        raise ValueError("Salidas inválidas.")
    expected_outputs = canonical_output_paths(Path(payload["ruta_candidatura"]))
    for key in OUTPUT_KEYS:
        if payload["salidas"][key] != expected_outputs[key].as_posix():
            raise ValueError(f"Salida no canónica: {key}")

    _assert_slot_map(payload["cv"], CV_FIELDS, "cv")
    _assert_slot_map(payload["carta"], CARTA_FIELDS, "carta")
    _assert_slot_map(payload["latex"], LATEX_FIELDS, "latex")
    _validate_experiences(payload["cv"])
    _validate_experiences(payload["latex"])
    for key in ("[NOMBRE]", "[TITULAR]", "[PERFIL PROFESIONAL]", "[PROPUESTA DE VALOR]", "[DESTINATARIO]", "[FECHA]", "[ASUNTO]", "[SALUDO]", "[APERTURA]", "[CIERRE]", "[DESPEDIDA]", "[FIRMA]"):
        section = payload["carta"] if key in CARTA_FIELDS else payload["cv"]
        if not section[key].strip():
            raise ValueError(f"Campo obligatorio vacío en carta/CV: {key}")
    return route


def validate_json_schema(payload: dict[str, Any], project_root: Path) -> None:
    schema_path = project_root / TEMPLATE_ROOT / "SCHEMA_DATOS_GENERACION_CANDIDATURA_1.0.json"
    if not schema_path.is_file():
        raise ValueError(f"Falta el esquema JSON: {schema_path}")
    schema = json.loads(schema_path.read_text(encoding="utf-8"))
    try:
        from jsonschema import Draft202012Validator
    except ImportError:
        # validate_payload performs the closed-contract checks when the optional
        # validator is not installed; the declared requirements enable the full
        # JSON Schema validation in a configured runtime.
        return
    errors = sorted(Draft202012Validator(schema).iter_errors(payload), key=lambda error: list(error.path))
    if errors:
        raise ValueError(f"JSON no válido según el esquema: {errors[0].message}")


VALID_STATES = {"en_preparacion", "pendiente_de_aprobacion", "detenida", "enviada", "rechazada", "aprobada", "duplicada", "fallida"}


def _frontmatter(path: Path) -> dict[str, str]:
    lines = path.read_text(encoding="utf-8").splitlines()
    if not lines or lines[0].strip() != "---":
        raise ValueError(f"Falta frontmatter en {path}.")
    try:
        end = lines.index("---", 1)
    except ValueError as exc:
        raise ValueError(f"Frontmatter sin cierre en {path}.") from exc
    result: dict[str, str] = {}
    for line in lines[1:end]:
        key, separator, value = line.partition(":")
        if separator:
            result[key.strip()] = value.strip()
    return result


def _markdown_row(line: str) -> list[str]:
    content = line.strip()
    if content.startswith("|"):
        content = content[1:]
    if content.endswith("|"):
        content = content[:-1]
    fields: list[str] = []
    current: list[str] = []
    wikilink_depth = 0
    position = 0
    while position < len(content):
        if content.startswith("[[", position):
            wikilink_depth += 1
            current.append("[[")
            position += 2
            continue
        if content.startswith("]]", position) and wikilink_depth:
            wikilink_depth -= 1
            current.append("]]" )
            position += 2
            continue
        character = content[position]
        if character == "|" and wikilink_depth == 0:
            fields.append("".join(current).strip())
            current = []
        else:
            current.append(character)
        position += 1
    fields.append("".join(current).strip())
    return fields


def validate_candidate_state(root: Path, route: Path, candidate_id: str) -> None:
    metadata = _frontmatter(route / "candidatura.md")
    declared_id = metadata.get("id") or metadata.get("id_candidatura")
    if declared_id != candidate_id:
        raise ValueError("candidatura.md debe declarar el mismo id que la candidatura.")
    state = metadata.get("estado")
    presented = metadata.get("presentada")
    if state not in VALID_STATES or presented not in {"true", "false"}:
        raise ValueError("candidatura.md debe declarar estado conocido y presentada: true|false.")
    tracking = root / "boveda-entrevista-profesional/busqueda-empleo/seguimiento/seguimiento-candidaturas.md"
    rows = tracking.read_text(encoding="utf-8").splitlines()
    headers = next((_markdown_row(line) for line in rows if line.startswith("| id_candidatura ")), None)
    if not headers or "estado" not in headers or "presentada" not in headers:
        raise ValueError("El seguimiento debe contener columnas id_candidatura, estado y presentada.")
    matching_rows = [_markdown_row(line) for line in rows if line.startswith(f"| {candidate_id} ")]
    if len(matching_rows) != 1 or len(matching_rows[0]) != len(headers):
        raise ValueError(f"No existe una fila única y completa para {candidate_id} en el seguimiento.")
    row = matching_rows[0]
    values = dict(zip(headers, row))
    if values.get("estado") != state or values.get("presentada") != presented:
        raise ValueError("El estado o presentada no coincide entre candidatura y seguimiento.")
    if presented == "true" or state in {"duplicada", "enviada", "rechazada"}:
        raise ValueError("La candidatura presentada o no regenerable no puede sobrescribirse.")
    if state not in {"en_preparacion", "pendiente_de_aprobacion", "detenida", "fallida", "aprobada"}:
        raise ValueError(f"El estado {state} no permite generar documentos.")


def _set_candidate_state(root: Path, route: Path, candidate_id: str, state: str) -> None:
    if state not in VALID_STATES:
        raise ValueError(f"Estado no permitido: {state}")
    candidature = route / "candidatura.md"
    text = candidature.read_text(encoding="utf-8")
    updated, count = re.subn(r"(?m)^estado:\s*.*$", f"estado: {state}", text, count=1)
    if count != 1:
        raise ValueError(f"No se pudo actualizar el estado en {candidature}.")
    candidature.write_text(updated, encoding="utf-8")

    tracking = root / "boveda-entrevista-profesional/busqueda-empleo/seguimiento/seguimiento-candidaturas.md"
    lines = tracking.read_text(encoding="utf-8").splitlines()
    headers = next((_markdown_row(line) for line in lines if line.startswith("| id_candidatura ")), None)
    if not headers or "estado" not in headers:
        raise ValueError("El seguimiento no permite actualizar el estado.")
    state_index = headers.index("estado")
    matching = [index for index, line in enumerate(lines) if line.startswith(f"| {candidate_id} ")]
    if len(matching) != 1:
        raise ValueError(f"No existe una fila única para {candidate_id} en el seguimiento.")
    fields = _markdown_row(lines[matching[0]])
    if len(fields) != len(headers):
        raise ValueError(f"La fila de seguimiento de {candidate_id} no es completa.")
    fields[state_index] = state
    lines[matching[0]] = "| " + " | ".join(fields) + " |"
    tracking.write_text("\n".join(lines) + "\n", encoding="utf-8")


def validate_resume_decision(route: Path) -> dict[str, Any]:
    review_path = route / "revision-generacion.json"
    decision_path = route / "revision-generacion-decision.json"
    if not review_path.is_file():
        raise ValueError(f"No existe una revisión de generación pendiente: {review_path}")
    review = json.loads(review_path.read_text(encoding="utf-8"))
    if review.get("estado") != "detenida_revision_humana":
        raise ValueError("La revisión de generación no está pendiente de decisión humana.")
    if not decision_path.is_file():
        raise ValueError(
            "Falta la decisión humana. Crea revision-generacion-decision.json con "
            "decision=corregir_y_reanudar o decision=aceptar_excepcion."
        )
    decision = json.loads(decision_path.read_text(encoding="utf-8"))
    if decision.get("decision") not in {"corregir_y_reanudar", "aceptar_excepcion"}:
        raise ValueError("La decisión humana debe ser corregir_y_reanudar o aceptar_excepcion.")
    if decision["decision"] == "aceptar_excepcion":
        if review.get("documento") != "cv.pdf":
            raise ValueError("Solo se admite una excepción de páginas para el CV.")
        pages = decision.get("paginas_aceptadas")
        if not isinstance(pages, int) or pages < review.get("paginas_reales", 1):
            raise ValueError("aceptar_excepcion requiere paginas_aceptadas igual o superior a las páginas reales.")
    return decision


def reopen_obsolete_page_review(root: Path, route: Path, candidate_id: str) -> bool:
    review_path = route / "revision-generacion.json"
    if not review_path.is_file():
        return False
    review = json.loads(review_path.read_text(encoding="utf-8"))
    if review.get("estado") != "detenida_revision_humana":
        return False
    limits = {"cv.pdf": 2, "carta-presentacion.pdf": 1}
    document = review.get("documento")
    actual_pages = review.get("paginas_reales")
    if document not in limits or not isinstance(actual_pages, int) or actual_pages > limits[document]:
        return False
    review["estado"] = "obsoleta_por_limite_actualizado"
    review["limite_actual"] = limits[document]
    review["reabierta_en"] = datetime.now(timezone.utc).isoformat()
    review_path.write_text(json.dumps(review, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    _set_candidate_state(root, route, candidate_id, "en_preparacion")
    return True


def _validate_photo(path: Path) -> None:
    if path.stat().st_size < 1024 or path.stat().st_size > 10 * 1024 * 1024:
        raise ValueError("La fotografía debe ocupar entre 1 KB y 10 MB.")
    try:
        with Image.open(path) as image:
            if image.format not in {"PNG", "JPEG"} or min(image.size) < 270:
                raise ValueError("La fotografía debe ser PNG/JPEG y medir al menos 270 px.")
            image.verify()
    except ValueError:
        raise
    except Exception as exc:
        raise ValueError("La fotografía no es una imagen válida.") from exc


def _replace_markers(text: str, values: dict[str, str]) -> str:
    return MARKER_RE.sub(lambda match: values.get(match.group(0), match.group(0)), text)


def build_latex(template: str, values: dict[str, str]) -> str:
    lines = template.splitlines(keepends=True)
    result: list[str] = []
    for line in lines:
        experience_numbers = {int(m.group(1)) for m in EXPERIENCE_RE.finditer(line)}
        if experience_numbers:
            number = next(iter(experience_numbers))
            header = values[f"[EXPERIENCIA {number} CABECERA]"]
            description = values[f"[EXPERIENCIA {number} DESCRIPCION]"]
            if not header and not description:
                continue
        normalized = line.strip().lower()
        if "[email]" in line.lower() and not any(values[key] for key in ("[EMAIL]", "[TELÉFONO]", "[LINKEDIN]")):
            continue
        if "\\section*{competencias y herramientas}" in normalized and not any(values[f"[COMPETENCIA {i}]"] for i in range(1, 5)):
            continue
        if "\\section*{formación}" in normalized and not any(values[f"[FORMACION {i}]"] for i in range(1, 4)):
            continue
        optional_markers = [marker for marker in _template_markers(line) if marker in values and not values[marker]]
        if optional_markers and all(marker in {f"[COMPETENCIA {i}]" for i in range(1, 5)} | {f"[FORMACION {i}]" for i in range(1, 4)} | {"[INFORMACION ADICIONAL]"} for marker in optional_markers):
            continue
        replaced = _replace_markers(line, values)
        if replaced.strip():
            result.append(replaced)
    return "".join(result)


def validate_latex(path: Path) -> str:
    text = path.read_text(encoding="utf-8")
    if any(marker in text for marker in LATEX_FIELDS):
        raise ValueError("LaTeX contiene marcadores sin sustituir.")
    if text.count("{") != text.count("}") or text.count("\\begin{") != text.count("\\end{"):
        raise ValueError("LaTeX tiene una estructura de llaves o entornos incompleta.")
    return "validado_estructuralmente"


def _iter_paragraphs(parent: Any) -> Iterable[Any]:
    yield from parent.paragraphs
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_paragraphs(cell)


def _rpr(run: Any) -> Any:
    return deepcopy(run._r.rPr) if run._r.rPr is not None else None


def _set_paragraph_text(paragraph: Any, text: str, style_overrides: list[tuple[int, int, Any]] | None = None) -> None:
    runs = list(paragraph.runs)
    styles: list[Any] = []
    for run in runs:
        styles.extend([_rpr(run)] * len(run.text or ""))
    while paragraph._p.r_lst:
        paragraph._p.remove(paragraph._p.r_lst[0])
    for position, char in enumerate(text):
        r = OxmlElement("w:r")
        index = min(len(styles) - 1, position) if styles else 0
        style = styles[index] if styles else None
        for start, end, override in style_overrides or []:
            if start <= position < end:
                style = override
                break
        if style is not None:
            r.append(deepcopy(style))
        t = OxmlElement("w:t")
        if char.isspace():
            t.set(qn("xml:space"), "preserve")
        t.text = char
        r.append(t)
        paragraph._p.append(r)


def _template_markers(text: str) -> list[str]:
    return MARKER_RE.findall(text)


def _validate_template_text(text: str, expected: tuple[str, ...], label: str, ignored: tuple[str, ...] = ()) -> None:
    markers = _template_markers(text)
    unknown = sorted(set(markers) - set(expected) - set(ignored))
    duplicates = sorted(marker for marker in set(markers) if markers.count(marker) > 1 and marker not in ignored)
    missing = sorted(set(expected) - set(markers))
    if unknown or duplicates or missing:
        raise ValueError(
            f"Contrato de marcadores inválido en {label}: "
            f"desconocidos={unknown}, duplicados={duplicates}, ausentes={missing}"
        )


def _template_paragraph_text(document: Document) -> str:
    paragraphs: list[str] = []
    for parent in [document, *[section.header for section in document.sections], *[section.footer for section in document.sections]]:
        paragraphs.extend("".join(run.text or "" for run in paragraph.runs) for paragraph in _iter_paragraphs(parent))
    return "\n".join(paragraphs)


def _replace_docx_paragraph(paragraph: Any, values: dict[str, str]) -> None:
    original = "".join(run.text or "" for run in paragraph.runs)
    if not MARKER_RE.search(original):
        return
    experience_styles: list[tuple[int, int, Any]] = []
    character_styles: list[Any] = []
    for run in paragraph.runs:
        character_styles.extend([_rpr(run)] * len(run.text or ""))
    for match in EXPERIENCE_RE.finditer(original):
        value = values.get(match.group(0), "")
        if not value:
            continue
        prefix = _replace_markers(original[:match.start()], values)
        style = character_styles[match.start()] if match.start() < len(character_styles) else None
        experience_styles.append((len(prefix), len(prefix) + len(value), style))
    replaced = _replace_markers(original, values)
    if "[EMAIL]" in original or "[TELÉFONO]" in original or "[LINKEDIN]" in original:
        parts = [values.get(key, "").strip() for key in ("[EMAIL]", "[TELÉFONO]", "[LINKEDIN]")]
        contact = " | ".join(part for part in parts if part)
        replaced = re.sub(r"\[EMAIL\] \| \[TELÉFONO\] \| \[LINKEDIN\]", contact, original)
        replaced = _replace_markers(replaced, values)
    if not replaced.strip():
        paragraph._p.getparent().remove(paragraph._p)
        return
    _set_paragraph_text(paragraph, replaced, experience_styles)


def build_docx(template_path: Path, output_path: Path, values: dict[str, str], photo_path: Path) -> None:
    document = Document(template_path)
    expected = CV_FIELDS if "CV" in template_path.name else CARTA_FIELDS
    required = {"[NOMBRE]", "[TITULAR]", "[PERFIL PROFESIONAL]", "[PROPUESTA DE VALOR]"} if expected is CV_FIELDS else {"[NOMBRE]", "[TITULAR]", "[EMAIL]", "[TELÉFONO]", "[LINKEDIN]", "[DESTINATARIO]", "[FECHA]", "[ASUNTO]", "[SALUDO]", "[APERTURA]", "[CIERRE]", "[DESPEDIDA]", "[FIRMA]"}
    _validate_template_text(_template_paragraph_text(document), expected, template_path.name)
    if not required.issubset(set(_template_markers(_template_paragraph_text(document)))):
        raise ValueError(f"Faltan marcadores obligatorios en {template_path.name}.")
    image_slots = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                count = len(cell._tc.xpath(".//wp:inline"))
                if count:
                    image_slots.append((cell, count))
    if len(image_slots) != 1 or image_slots[0][1] != 1 or len(document.inline_shapes) != 1:
        raise ValueError("La plantilla no contiene exactamente un slot de fotografía.")
    _replace_image(document, photo_path)
    for parent in [document, *[section.header for section in document.sections], *[section.footer for section in document.sections]]:
        for paragraph in list(_iter_paragraphs(parent)):
            text = "".join(run.text or "" for run in paragraph.runs)
            normalized = text.strip().upper()
            if normalized == "COMPETENCIAS Y HERRAMIENTAS" and not any(values[f"[COMPETENCIA {i}]"] for i in range(1, 5)):
                paragraph._p.getparent().remove(paragraph._p)
                continue
            if normalized in {"FORMACIÓN", "FORMACION"} and not any(values[f"[FORMACION {i}]"] for i in range(1, 4)):
                paragraph._p.getparent().remove(paragraph._p)
                continue
            if normalized == "INFORMACIÓN ADICIONAL" or normalized == "INFORMACION ADICIONAL":
                if not values["[INFORMACION ADICIONAL]"]:
                    paragraph._p.getparent().remove(paragraph._p)
                    continue
            experience = EXPERIENCE_RE.search(text)
            if experience:
                number = int(experience.group(1))
                if not values[f"[EXPERIENCIA {number} CABECERA]"] and not values[f"[EXPERIENCIA {number} DESCRIPCION]"]:
                    paragraph._p.getparent().remove(paragraph._p)
                    continue
            _replace_docx_paragraph(paragraph, values)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def _replace_image(document: Document, photo_path: Path) -> None:
    shape = document.inline_shapes[0]
    with Image.open(photo_path) as image:
        image = image.convert("RGB")
        side = min(image.size)
        left = (image.width - side) // 2
        top = (image.height - side) // 2
        image = image.crop((left, top, left + side, top + side))
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
            image.save(temp.name, format="PNG")
            png = Path(temp.name).read_bytes()
        Path(temp.name).unlink(missing_ok=True)
    blip = shape._inline.xpath(".//a:blip")[0]
    relation_id = blip.get(qn("r:embed"))
    image_part = document.part.related_parts[relation_id]
    image_part._blob = png


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _load_env() -> dict[str, str]:
    path = SCRIPT_DIR / ".env"
    if not path.is_file():
        raise ValueError(f"Falta la configuración: {path}")
    values: dict[str, str] = {}
    for line in path.read_text(encoding="utf-8").splitlines():
        line = line.strip()
        if not line or line.startswith("#"):
            continue
        key, separator, value = line.partition("=")
        if separator:
            values[key.strip()] = value.strip().strip('"')
    for key in ("RUTA_PROYECTO", "SOFFICE_PATH"):
        if not values.get(key):
            raise ValueError(f"Falta {key} en {path}")
    if not Path(values["RUTA_PROYECTO"]).is_absolute():
        raise ValueError("RUTA_PROYECTO debe ser una ruta absoluta.")
    soffice = Path(values["SOFFICE_PATH"])
    if not soffice.is_absolute() or not soffice.is_file():
        raise ValueError("SOFFICE_PATH debe apuntar a un ejecutable existente.")
    return values


def convert_docx_to_pdf(docx_path: Path, out_dir: Path, soffice_path: Path, conversion_root: Path | None = None) -> Path:
    out_dir.mkdir(parents=True, exist_ok=True)
    expected = out_dir / f"{docx_path.stem}.pdf"
    expected.unlink(missing_ok=True)
    staging = (conversion_root / ".tmp/job-up-lo" / uuid.uuid4().hex) if conversion_root is not None else (out_dir / f"lo-{uuid.uuid4().hex}")
    staging.mkdir(parents=True, exist_ok=False)
    staged_docx = staging / docx_path.name
    shutil.copy2(docx_path, staged_docx)
    staged_pdf = staging / f"{staged_docx.stem}.pdf"
    profile = staging / "profile"
    profile.mkdir()
    uri = profile.resolve().as_uri()
    command = [str(soffice_path), "--headless", "--nologo", "--nodefault", "--nofirststartwizard", "--norestore", f"-env:UserInstallation={uri}", "--convert-to", "pdf", "--outdir", str(staging), str(staged_docx)]
    try:
        process = subprocess.Popen(command, cwd=staging, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        try:
            stdout, stderr = process.communicate(timeout=DOCX_TO_PDF_TIMEOUT)
        except subprocess.TimeoutExpired as exc:
            if os.name == "nt":
                subprocess.run(["taskkill", "/PID", str(process.pid), "/T", "/F"], capture_output=True, text=True)
            else:
                process.kill()
            process.wait()
            _assert_process_tree_gone(process.pid)
            timeout_stderr = getattr(exc, "stderr", None)
            raise RuntimeError(f"LibreOffice superó {DOCX_TO_PDF_TIMEOUT} segundos: {timeout_stderr or exc}") from exc
        if process.returncode != 0 or not staged_pdf.is_file() or staged_pdf.stat().st_size == 0:
            raise RuntimeError(f"Conversión PDF fallida ({process.returncode}): {stderr or stdout}")
        shutil.copy2(staged_pdf, expected)
        return expected
    finally:
        shutil.rmtree(staging, ignore_errors=True)


class PdfValidationError(ValueError):
    def __init__(self, path: Path, actual_pages: int, max_pages: int) -> None:
        self.path = path
        self.actual_pages = actual_pages
        self.max_pages = max_pages
        super().__init__(
            f"{path.name}: tiene {actual_pages} páginas; el máximo permitido es {max_pages}. "
            "Revisa el documento y decide si corregirlo o aceptar expresamente la excepción."
        )


class HumanReviewRequired(RuntimeError):
    resultado = "revision_humana_requerida"
    codigo_error = "REVISION_HUMANA_REQUERIDA"


def validate_pdf(path: Path, require_image: bool = False, max_pages: int = 1) -> None:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    actual_pages = len(reader.pages)
    if actual_pages > max_pages:
        raise PdfValidationError(path, actual_pages, max_pages)
    text = "\n".join(page.extract_text() or "" for page in reader.pages)
    if "[" in text and "]" in text:
        raise ValueError(f"PDF contiene marcadores: {path}")
    if require_image and not any(getattr(page, "images", ()) for page in reader.pages):
        raise ValueError(f"PDF sin fotografía embebida: {path}")


def preserve_generation_review(
    root: Path,
    route: Path,
    execution_dir: Path,
    execution_id: str,
    document: str,
    actual_pages: int,
    max_pages: int,
) -> dict[str, Any]:
    review_dir = route / "revisiones-generacion" / execution_id
    review_dir.mkdir(parents=True, exist_ok=False)
    artifacts: dict[str, str] = {}
    for name in ("cv.docx", "cv.pdf", "carta-presentacion.docx", "carta-presentacion.pdf"):
        source = execution_dir / name
        if source.is_file():
            destination = review_dir / name
            shutil.copy2(source, destination)
            artifacts[name.replace("-", "_").replace(".", "_").rstrip("_")] = str(destination)
    review = {
        "schema_version": "1.0",
        "estado": "detenida_revision_humana",
        "execution_id": execution_id,
        "documento": document,
        "paginas_reales": actual_pages,
        "paginas_maximas": max_pages,
        "artefactos": artifacts,
        "siguiente_accion": (
            "Mostrar al usuario la ruta del PDF conservado y esperar una "
            "respuesta explícita de sí o no antes de reanudar."
        ),
        "decision_admitida": ["corregir_y_reanudar", "aceptar_excepcion"],
    }
    review_path = route / "revision-generacion.json"
    review_path.write_text(json.dumps(review, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    candidate_id = _frontmatter(route / "candidatura.md").get("id") or route.name.split("-", 3)[0]
    _set_candidate_state(root, route, candidate_id, "detenida")
    return review


def write_error_record(root: Path, execution_id: str, input_path: str, exc: Exception, phase: str, candidate_id: str | None = None) -> Path | None:
    manifest = getattr(exc, "manifest", {})
    files = manifest.get("files", []) if isinstance(manifest, dict) else []
    published = manifest.get("publicados", []) if isinstance(manifest, dict) else []
    restored = manifest.get("documentos_restaurados", []) if isinstance(manifest, dict) else []
    unresolved = [entry.get("destination", "") for entry in files if entry.get("destination") not in set(restored) and entry.get("destination") not in set(published)]
    record = {
        "schema_version": "1.0",
        "fecha": datetime.now(timezone.utc).isoformat(),
        "resultado": getattr(exc, "resultado", "fallido"),
        "execution_id": execution_id,
        "id_candidatura": candidate_id,
        "entrada_recibida": input_path,
        "fase": phase,
        "codigo_error": getattr(exc, "codigo_error", "GENERACION_FALLIDA"),
        "campo_o_ruta": str(getattr(exc, "path", "")) or None,
        "mensaje": str(exc) or exc.__class__.__name__,
        "documentos_publicados": published,
        "documentos_restaurados": restored,
        "documentos_sin_publicar_o_restaurar": unresolved,
    }
    target_dir = root / ERROR_ROOT
    try:
        target_dir.mkdir(parents=True, exist_ok=True)
        target = target_dir / f"error-{datetime.now().strftime('%Y%m%d-%H%M%S')}-{execution_id}.json"
        target.write_text(json.dumps(record, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
        return target
    except OSError:
        return None


def _write_json_atomic(path: Path, value: dict[str, Any]) -> None:
    temporary = path.with_name(f".{path.name}.{uuid.uuid4().hex}.tmp")
    temporary.write_text(json.dumps(value, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    os.replace(temporary, path)


class PublicationError(OSError):
    def __init__(self, message: str, manifest: dict[str, Any]) -> None:
        super().__init__(message)
        self.manifest = manifest


class CandidateLock:
    def __init__(self, path: Path) -> None:
        self.path = path
        self._handle: int | None = None

    def __enter__(self) -> "CandidateLock":
        self.path.parent.mkdir(parents=True, exist_ok=True)
        if self.path.exists():
            try:
                owner = json.loads(self.path.read_text(encoding="utf-8"))
                pid = int(owner.get("pid", ""))
            except (OSError, ValueError, TypeError, json.JSONDecodeError) as exc:
                raise RuntimeError(f"Bloqueo ilegible; requiere decisión humana: {self.path}") from exc
            if not _pid_is_running(pid):
                self.path.unlink()
            else:
                raise RuntimeError(f"Ya existe una ejecución activa para la candidatura (pid {pid}).")
        try:
            self._handle = os.open(self.path, os.O_CREAT | os.O_EXCL | os.O_WRONLY)
            os.write(self._handle, json.dumps({"pid": os.getpid(), "fecha": datetime.now(timezone.utc).isoformat()}).encode("utf-8"))
        except FileExistsError as exc:
            raise RuntimeError("Ya existe una ejecución activa para la candidatura.") from exc
        return self

    def __exit__(self, exc_type: Any, exc: Any, traceback: Any) -> None:
        if self._handle is not None:
            os.close(self._handle)
            self._handle = None
            self.path.unlink(missing_ok=True)


def _pid_is_running(pid: int) -> bool:
    if os.name == "nt":
        result = subprocess.run(
            ["tasklist", "/FI", f"PID eq {pid}", "/FO", "CSV", "/NH"],
            capture_output=True,
            text=True,
            check=False,
        )
        return result.returncode == 0 and str(pid) in result.stdout
    try:
        os.kill(pid, 0)
    except OSError:
        return False
    return True


def _assert_process_tree_gone(pid: int) -> None:
    if os.name != "nt":
        return
    if _pid_is_running(pid):
        raise RuntimeError(f"El proceso LibreOffice {pid} sigue activo tras la terminación.")
    result = subprocess.run(
        ["powershell", "-NoProfile", "-Command", f"(Get-CimInstance Win32_Process -Filter 'ParentProcessId={pid}').ProcessId"],
        capture_output=True,
        text=True,
        check=False,
    )
    if result.returncode != 0:
        raise RuntimeError("No se pudo comprobar la ausencia de procesos descendientes de LibreOffice.")
    if result.stdout.strip():
        raise RuntimeError(f"Quedan procesos descendientes de LibreOffice: {result.stdout.strip()}")


def _manifest_path(execution_dir: Path) -> Path:
    return execution_dir / "manifest.json"


def _restore_manifest(manifest_path: Path, root: Path) -> dict[str, Any]:
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    backup_dir = manifest_path.parent / "backups"
    restored: list[str] = []
    missing: list[str] = []
    for entry in manifest["files"]:
        destination = root / Path(entry["destination"])
        backup = backup_dir / entry["backup_name"]
        if entry["had_original"]:
            if not backup.is_file():
                missing.append(entry["destination"])
                continue
            os.replace(backup, destination)
        elif destination.exists():
            destination.unlink()
        restored.append(entry["destination"])
    if missing:
        manifest["phase"] = "restauracion_incompleta"
        manifest["documentos_restaurados"] = restored
        manifest["documentos_sin_restaurar"] = missing
        _write_json_atomic(manifest_path, manifest)
        raise RuntimeError(f"No se pudo restaurar: {', '.join(missing)}")
    manifest["phase"] = "restaurado"
    manifest["documentos_restaurados"] = restored
    _write_json_atomic(manifest_path, manifest)
    return manifest


def recover_pending_publications(root: Path, candidate_name: str) -> list[Path]:
    candidate_temp = root / TEMP_ROOT / candidate_name
    if not candidate_temp.is_dir():
        return []
    recovered: list[Path] = []
    for execution_dir in sorted(path for path in candidate_temp.iterdir() if path.is_dir()):
        manifest_path = _manifest_path(execution_dir)
        if not manifest_path.is_file():
            continue
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        if manifest.get("phase") in {"publicando", "restaurando", "restauracion_incompleta"}:
            _restore_manifest(manifest_path, root)
            recovered.append(manifest_path)
    return recovered


def publish_transaction(
    root: Path,
    execution_dir: Path,
    generated: dict[str, Path],
    destinations: dict[str, Path],
    execution_id: str,
    max_pdf_pages: dict[str, int] | None = None,
) -> dict[str, Any]:
    backup_dir = execution_dir / "backups"
    backup_dir.mkdir(parents=True, exist_ok=True)
    files: list[dict[str, Any]] = []
    for key in OUTPUT_KEYS:
        source = generated[key]
        destination = destinations[key]
        destination.parent.mkdir(parents=True, exist_ok=True)
        had_original = destination.is_file()
        backup_name = f"{key}.bak"
        if had_original:
            shutil.copy2(destination, backup_dir / backup_name)
        files.append({
            "key": key,
            "source": source.name,
            "destination": destination.relative_to(root).as_posix(),
            "backup_name": backup_name,
            "had_original": had_original,
            "sha256": sha256(source),
        })
    manifest = {"schema_version": "1.0", "execution_id": execution_id, "phase": "publicando", "files": files, "publicados": []}
    manifest_path = _manifest_path(execution_dir)
    _write_json_atomic(manifest_path, manifest)
    try:
        for entry in files:
            source = execution_dir / entry["source"]
            destination = root / Path(entry["destination"])
            os.replace(source, destination)
            if sha256(destination) != entry["sha256"]:
                raise RuntimeError(f"Hash publicado incorrecto: {destination}")
            manifest["publicados"].append(entry["destination"])
            _write_json_atomic(manifest_path, manifest)
        if all(path.suffix.lower() in {".docx", ".pdf", ".tex"} for path in destinations.values()):
            validate_published_artifacts(destinations, max_pdf_pages)
        manifest["phase"] = "completado"
        _write_json_atomic(manifest_path, manifest)
        return manifest
    except Exception as exc:
        manifest["phase"] = "restaurando"
        _write_json_atomic(manifest_path, manifest)
        try:
            manifest = _restore_manifest(manifest_path, root)
        except Exception as restoration_error:
            setattr(restoration_error, "manifest", manifest)
            raise
        raise PublicationError(str(exc), manifest) from exc


def cleanup_execution_directory(path: Path) -> None:
    if path.is_dir():
        shutil.rmtree(path)


def validate_published_artifacts(destinations: dict[str, Path], max_pdf_pages: dict[str, int] | None = None) -> None:
    max_pdf_pages = max_pdf_pages or {}
    for key, path in destinations.items():
        if not path.is_file() or path.stat().st_size == 0:
            raise ValueError(f"Artefacto publicado ausente o vacío: {key}")
        if key.endswith("docx"):
            Document(path)
        elif key.endswith("pdf"):
            validate_pdf(path, require_image=True, max_pages=max_pdf_pages.get(key, 1))
        elif key == "cv_tex":
            text = path.read_text(encoding="utf-8")
            if any(marker in text for marker in LATEX_FIELDS):
                raise ValueError("cv.tex publicado contiene marcadores sin sustituir.")


def _contains_placeholder(value: Any) -> bool:
    if isinstance(value, dict):
        return any(_contains_placeholder(item) for item in value.values())
    if isinstance(value, list):
        return any(_contains_placeholder(item) for item in value)
    return isinstance(value, str) and bool(re.search(r"{{[^}]+}}", value))


PRIVATE_DATA_FIELDS = (
    "nombre",
    "apellido_1",
    "apellido_2",
    "email",
    "telefono",
    "linkedin",
    "ubicacion",
    "fotografia",
)


def validate_private_data_manifest(payload: dict[str, Any]) -> None:
    """Valida la autorización explícita y la materialización de datos privados."""
    manifest = payload.get("control", {}).get("datos_privados")
    if not isinstance(manifest, dict):
        raise ValueError("Falta el contrato de autorización de datos privados.")
    authorization = manifest.get("autorizacion")
    if not isinstance(authorization, dict) or set(authorization) != set(PRIVATE_DATA_FIELDS):
        raise ValueError("La autorización de datos privados debe resolver todos los campos.")
    if any(value not in {"incluir", "omitir"} for value in authorization.values()):
        raise ValueError("La autorización de datos privados no puede quedar pendiente.")
    if not isinstance(manifest.get("fecha_decision"), str) or not manifest["fecha_decision"].strip():
        raise ValueError("Falta la fecha de decisión de datos privados.")
    if manifest.get("decidido_por") != "persona_responsable":
        raise ValueError("La decisión de datos privados debe proceder de la persona responsable.")

    encabezado = payload.get("contenido_cv", {}).get("encabezado", {})
    nombre = encabezado.get("nombre_completo")
    if not isinstance(nombre, dict):
        raise ValueError("Falta el nombre del encabezado.")
    nombre_refs = {
        referencia
        for origen in nombre.get("trazabilidad", {}).get("origen_factual", [])
        if isinstance(origen, dict) and origen.get("fuente") == "datos-privados-candidatura"
        for referencia in origen.get("refs", [])
    }
    name_field_by_ref = {"Nombre": "nombre", "Apellido 1": "apellido_1", "Apellido 2": "apellido_2"}
    expected_name_refs = {
        reference for reference, field in name_field_by_ref.items() if authorization[field] == "incluir"
    }
    if nombre_refs != expected_name_refs:
        raise ValueError("El nombre visible no coincide con la autorización de nombre y apellidos.")

    contact = encabezado.get("contacto", [])
    if not isinstance(contact, list):
        raise ValueError("El contacto del encabezado debe ser una lista.")
    contact_types = {item.get("tipo") for item in contact if isinstance(item, dict)}
    authorized_contact = {field for field in ("email", "telefono", "linkedin", "ubicacion") if authorization[field] == "incluir"}
    if contact_types != authorized_contact:
        raise ValueError("El contacto visible no coincide con la autorización de datos privados.")
    if any(not isinstance(item, dict) or not item.get("texto", "").strip() for item in contact):
        raise ValueError("El contacto autorizado no puede estar vacío.")
    if authorization["fotografia"] != "incluir":
        raise ValueError("La excepción de fotografía sin imagen requiere una decisión contractual específica.")


def validate_composition_payload(payload: dict[str, Any]) -> str:
    """Valida la frontera mínima que necesita el compositor CV-only."""
    if not isinstance(payload, dict):
        raise ValueError("datos-generacion.json debe contener un objeto JSON.")
    if payload.get("schema_id") != "datos-generacion-cv" or payload.get("schema_version") != "1.2":
        raise ValueError("El compositor requiere el contrato datos-generacion-cv 1.2.")
    if payload.get("tipo") != "datos_generacion_cv":
        raise ValueError("El tipo debe ser datos_generacion_cv.")
    candidatura = payload.get("candidatura")
    candidate_id = candidatura.get("id") if isinstance(candidatura, dict) else None
    if not isinstance(candidate_id, str) or not re.fullmatch(r"CAND-\d{4}-\d{3}", candidate_id):
        raise ValueError("Identificador de candidatura inválido.")
    if not isinstance(payload.get("contenido_cv"), dict):
        raise ValueError("Falta contenido_cv.")
    validate_private_data_manifest(payload)
    validations = payload.get("control", {}).get("validaciones", {})
    if not validations or any(value is not True for value in validations.values()):
        raise ValueError("El contenido no tiene todas las validaciones aprobadas.")
    if _contains_placeholder(payload):
        raise ValueError("El JSON contiene placeholders sin resolver.")
    return candidate_id


def _candidate_route_from_json(root: Path, json_path: Path, candidate_id: str) -> Path:
    route = json_path.parent.resolve()
    _assert_inside(route, (root / CANDIDATURE_ROOT).resolve(), "carpeta de candidatura")
    if not route.name.startswith(candidate_id):
        raise ValueError("La carpeta de candidatura no coincide con candidatura.id.")
    return route


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("json_path")
    parser.add_argument(
        "--reanudar",
        action="store_true",
        help="reanuda una generación detenida solo con una decisión humana registrada",
    )
    args = parser.parse_args(argv)
    execution = uuid.uuid4().hex[:12]
    phase = "carga_json"
    root: Path | None = None
    candidate_id: str | None = None
    execution_dir: Path | None = None
    candidate_lock: CandidateLock | None = None
    try:
        env = _load_env()
        root = Path(env["RUTA_PROYECTO"]).resolve()
        phase = "configuracion"
        json_path = resolve_input_json(root, args.json_path)
        payload = json.loads(json_path.read_text(encoding="utf-8"))
        candidate_value = payload.get("candidatura", {}).get("id") if isinstance(payload, dict) else None
        candidate_id = candidate_value if isinstance(candidate_value, str) and re.fullmatch(r"CAND-\d{4}-\d{3}", candidate_value) else None
        phase = "validacion_entradas"
        candidate_id = validate_composition_payload(payload)
        route = _candidate_route_from_json(root, json_path, candidate_id)
        if args.reanudar:
            raise ValueError("La reanudación del flujo histórico no forma parte del contrato CV-only 1.2.")
        template = (root / TEMPLATE_ROOT / "TEMPLATE_CV_FORMATO.docx").resolve()
        photo = (root / "boveda-entrevista-profesional/busqueda-empleo/foto-perfil.png").resolve()
        if not template.is_file():
            raise ValueError(f"Falta la plantilla canónica: {template}")
        if not photo.is_file():
            raise ValueError(f"Falta la fotografía canónica: {photo}")
        _validate_photo(photo)
        model = construir_modelo_cv(payload)
        candidate_lock = CandidateLock(root / TEMP_ROOT / route.name / ".lock")
        candidate_lock.__enter__()
        recover_pending_publications(root, route.name)
        execution_dir = root / TEMP_ROOT / route.name / execution
        execution_dir.mkdir(parents=True, exist_ok=False)
        phase = "generacion_docx"
        cv_docx = execution_dir / "cv.docx"
        renderizar_docx(model, template, cv_docx, photo)
        phase = "conversion_pdf"
        cv_pdf = convert_docx_to_pdf(cv_docx, execution_dir, Path(env["SOFFICE_PATH"]), root)
        phase = "validacion_pdf"
        try:
            max_cv_pages = 2
            validate_pdf(cv_pdf, require_image=True, max_pages=max_cv_pages)
        except PdfValidationError as exc:
            review = preserve_generation_review(
                root,
                route,
                execution_dir,
                execution,
                exc.path.name,
                exc.actual_pages,
                exc.max_pages,
            )
            review_path = route / "revision-generacion.json"
            review_exc = HumanReviewRequired(str(exc))
            review_exc.path = exc.path
            review_exc.actual_pages = exc.actual_pages
            review_exc.max_pages = exc.max_pages
            review_exc.review_path = review_path
            raise review_exc from exc
        phase = "generacion_latex"
        tex = renderizar_latex(model)
        (execution_dir / "cv.tex").write_text(tex, encoding="utf-8", newline="\n")
        latex_status = validate_latex(execution_dir / "cv.tex")
        generated = {"cv_docx": cv_docx, "cv_pdf": cv_pdf, "cv_tex": execution_dir / "cv.tex"}
        destinations = canonical_output_paths(route)
        phase = "validacion_salidas"
        for key, source in generated.items():
            if not source.is_file() or source.stat().st_size == 0:
                raise ValueError(f"Artefacto temporal ausente o vacío: {key}")
        phase = "publicacion"
        publish_transaction(
            root,
            execution_dir,
            generated,
            destinations,
            execution,
            {"cv_pdf": max_cv_pages},
        )
        manifest_summary = {
            "schema_version": "1.0",
            "execution_id": execution,
            "id_candidatura": candidate_id,
            "fotografia": "incluida",
            "artefactos": list(OUTPUT_KEYS),
        }
        _write_json_atomic(route / "manifest-generacion-cv.json", manifest_summary)
        cleanup_execution_directory(execution_dir)
        execution_dir = None
        print(json.dumps({"resultado": "completado", "execution_id": execution, "latex": latex_status, "artefactos": list(destinations)}, ensure_ascii=False))
        return 0
    except Exception as exc:
        record = write_error_record(root, execution, args.json_path, exc, phase, candidate_id) if root is not None else None
        if record is not None:
            prefix = "REVISION_HUMANA_REQUERIDA" if getattr(exc, "resultado", "") == "revision_humana_requerida" else "GENERACION_FALLIDA"
            print(f"{prefix}: {exc}; registro={record}", file=sys.stderr)
        else:
            print(f"GENERACION_FALLIDA: {exc}; no se pudo escribir el registro de error", file=sys.stderr)
        return 2 if getattr(exc, "resultado", "") == "revision_humana_requerida" else 1
    finally:
        if execution_dir is not None and execution_dir.exists():
            manifest = execution_dir / "manifest.json"
            keep = False
            if manifest.is_file():
                try:
                    keep = json.loads(manifest.read_text(encoding="utf-8")).get("phase") in {"publicando", "restaurando", "restauracion_incompleta"}
                except (OSError, json.JSONDecodeError):
                    keep = True
            if not keep:
                cleanup_execution_directory(execution_dir)
        if candidate_lock is not None:
            candidate_lock.__exit__(None, None, None)


if __name__ == "__main__":
    raise SystemExit(main())
