"""Composición determinista de CV desde ``contenido_cv`` 1.2."""

from __future__ import annotations

import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image

from cabecera_candidatura import construir_cabecera_candidatura


@dataclass(frozen=True)
class RenderTexto:
    texto: str
    orden: int
    tipo: str


@dataclass(frozen=True)
class RenderBloque:
    id_bloque: str
    tipo: str
    orden: int
    cabecera: tuple[RenderTexto, ...]
    unidades: tuple[RenderTexto, ...]


@dataclass(frozen=True)
class RenderSeccion:
    id_seccion: str
    tipo: str
    titulo_visible: str | None
    orden: int
    bloques: tuple[RenderBloque, ...]


@dataclass(frozen=True)
class RenderEncabezado:
    nombre: str
    unidades: tuple[RenderTexto, ...]
    contacto: tuple[RenderTexto, ...]


@dataclass(frozen=True)
class RenderCV:
    encabezado: RenderEncabezado
    secciones: tuple[RenderSeccion, ...]


def _orden(item: dict[str, Any]) -> int:
    return int(item.get("orden", 0))


def _texto(item: dict[str, Any], discriminator: str) -> RenderTexto:
    return RenderTexto(str(item.get("texto", "")), _orden(item), str(item.get(discriminator, "linea")))


def construir_modelo_cv(documento: dict[str, Any]) -> RenderCV:
    """Construye el árbol visible usando exclusivamente ``contenido_cv``."""
    contenido = documento["contenido_cv"]
    canonical_header = construir_cabecera_candidatura(documento, validar_privacidad=False)
    encabezado = contenido["encabezado"]
    unidades = tuple(_texto(item, "tipo") for item in sorted(encabezado.get("unidades", []), key=_orden))
    contacto = tuple(_texto(item, "tipo") for item in sorted(encabezado.get("contacto", []), key=_orden))
    secciones: list[RenderSeccion] = []
    for section in sorted(contenido.get("secciones", []), key=_orden):
        blocks: list[RenderBloque] = []
        for block in sorted(section.get("bloques", []), key=_orden):
            blocks.append(RenderBloque(
                id_bloque=str(block["id_bloque"]),
                tipo=str(block["tipo"]),
                orden=_orden(block),
                cabecera=tuple(_texto(item, "rol") for item in sorted(block.get("cabecera", []), key=_orden)),
                unidades=tuple(_texto(item, "tipo") for item in sorted(block.get("unidades", []), key=_orden)),
            ))
        secciones.append(RenderSeccion(
            id_seccion=str(section["id_seccion"]),
            tipo=str(section["tipo"]),
            titulo_visible=section.get("titulo_visible"),
            orden=_orden(section),
            bloques=tuple(blocks),
        ))
    return RenderCV(
        RenderEncabezado(canonical_header.nombre, unidades, contacto),
        tuple(secciones),
    )


_LATEX_REPLACEMENTS = {
    "\\": r"\textbackslash{}", "&": r"\&", "%": r"\%", "$": r"\$",
    "#": r"\#", "_": r"\_", "{": r"\{", "}": r"\}",
    "~": r"\textasciitilde{}", "^": r"\textasciicircum{}",
}


def _latex(text: str) -> str:
    return "".join(_LATEX_REPLACEMENTS.get(character, character) for character in text)


def renderizar_latex(modelo: RenderCV) -> str:
    """Renderiza LaTeX dinámico desde el modelo común."""
    lines = [
        r"\documentclass[11pt,a4paper]{article}", r"\usepackage[utf8]{inputenc}",
        r"\usepackage[T1]{fontenc}", r"\usepackage[margin=1.6cm]{geometry}",
        r"\usepackage{xcolor}", r"\usepackage{enumitem}", r"\usepackage{hyperref}",
        r"\setlength{\parindent}{0pt}", r"\setlength{\parskip}{4pt}",
        r"\definecolor{primary}{HTML}{1F2937}", r"\definecolor{secondary}{HTML}{5B6573}",
        r"\begin{document}",
        rf"{{\color{{primary}}\LARGE\textbf{{{_latex(modelo.encabezado.nombre)}}}}}\\",
    ]
    for unit in modelo.encabezado.unidades:
        lines.append(rf"{{\color{{secondary}}\large {_latex(unit.texto)}}}\\")
    if modelo.encabezado.contacto:
        lines.append(r" \textbar{} ".join(_latex(item.texto) for item in modelo.encabezado.contacto))
    for section in modelo.secciones:
        if section.titulo_visible:
            lines.extend(["", rf"\section*{{{_latex(section.titulo_visible)}}}"])
        for block in section.bloques:
            if block.cabecera:
                lines.append(r" \textbar{} ".join(rf"\textbf{{{_latex(item.texto)}}}" for item in block.cabecera))
            bullets = [item for item in block.unidades if item.tipo == "bullet"]
            others = [item for item in block.unidades if item.tipo != "bullet"]
            for item in others:
                lines.append(_latex(item.texto) + (r"\\" if item.tipo == "linea" else ""))
            if bullets:
                lines.append(r"\begin{itemize}[leftmargin=*,nosep]")
                lines.extend(rf"\item {_latex(item.texto)}" for item in bullets)
                lines.append(r"\end{itemize}")
    lines.extend(["", r"\end{document}", ""])
    return "\n".join(lines)


def _replace_text_in_paragraph(paragraph: Any, replacements: dict[str, str]) -> None:
    original = "".join(run.text or "" for run in paragraph.runs)
    updated = original
    for marker, value in replacements.items():
        updated = updated.replace(marker, value)
    if updated == original:
        return
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = updated
    else:
        paragraph.add_run(updated)


def _replace_photo(document: Document, photo: Path) -> None:
    if len(document.inline_shapes) != 1:
        raise ValueError("La plantilla debe contener un único espacio de fotografía.")
    shape = document.inline_shapes[0]
    with Image.open(photo) as image:
        image = image.convert("RGB")
        side = min(image.size)
        left = (image.width - side) // 2
        top = (image.height - side) // 2
        image = image.crop((left, top, left + side, top + side))
        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp:
            temp_path = Path(temp.name)
            image.save(temp_path, format="PNG")
    try:
        blip = shape._inline.xpath(".//a:blip")[0]
        image_part = document.part.related_parts[blip.get(qn("r:embed"))]
        image_part._blob = temp_path.read_bytes()
    finally:
        temp_path.unlink(missing_ok=True)


def renderizar_docx(modelo: RenderCV, plantilla: Path, destino: Path, fotografia: Path) -> None:
    """Renderiza un DOCX dinámico conservando la cabecera visual canónica."""
    document = Document(plantilla)
    if not fotografia.is_file():
        raise ValueError(f"Falta la fotografía canónica: {fotografia}")
    replacements = {
        "[NOMBRE]": modelo.encabezado.nombre,
        "[TITULAR]": " | ".join(item.texto for item in modelo.encabezado.unidades),
        "[EMAIL] | [TELÉFONO] | [LINKEDIN]": " | ".join(item.texto for item in modelo.encabezado.contacto),
    }
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_text_in_paragraph(paragraph, replacements)
    _replace_photo(document, fotografia)
    for paragraph in list(document.paragraphs):
        paragraph._element.getparent().remove(paragraph._element)
    for section in modelo.secciones:
        if section.titulo_visible:
            document.add_heading(section.titulo_visible, level=1)
        for block in section.bloques:
            if block.cabecera:
                paragraph = document.add_paragraph()
                paragraph.add_run(" | ".join(item.texto for item in block.cabecera)).bold = True
            for item in block.unidades:
                paragraph = document.add_paragraph(style="List Bullet" if item.tipo == "bullet" else None)
                paragraph.add_run(item.texto)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if item.tipo == "parrafo" else WD_ALIGN_PARAGRAPH.LEFT
    destino.parent.mkdir(parents=True, exist_ok=True)
    document.save(destino)
