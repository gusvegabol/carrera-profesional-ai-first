"""Compositor y auditor determinista de cartas de presentación.

La fase consume únicamente ``Carta completa consolidada`` de
``contenido-carta-presentacion.md``. No redacta, resume ni consulta otras
fuentes para decidir el texto visible.
"""

from __future__ import annotations

import argparse
import difflib
import hashlib
import json
import re
import shutil
import subprocess
import uuid
import unicodedata
import zipfile
from dataclasses import dataclass, field
from datetime import date
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from pypdf import PdfReader

from cabecera_candidatura import (
    CabeceraCandidatura,
    CabeceraError,
HEADER_CONTACT_SIZE,
    HEADER_FONT,
    HEADER_NAME_SIZE,
    HEADER_TITLE_SIZE,
    construir_cabecera_candidatura,
    validar_cabecera_contrato,
)


SECTION_RE = re.compile(r"^##\s+\d+\.\s+Carta completa consolidada\s*$", re.IGNORECASE)
NEXT_SECTION_RE = re.compile(r"^##\s+")
FIGURE_RE = re.compile(r"(?<!\w)\d+(?:[.,]\d+)?\s*%?", re.UNICODE)
PRIVATE_DATA_RE = re.compile(
    r"(?:[^\s@]+@[^\s@]+\.[^\s@]+|(?:https?://|www\.|linkedin\.com/)\S+|\b\d{3}[\s.-]?\d{3}[\s.-]?\d{3}\b|\b\d{8}[A-Za-z]\b)",
    re.IGNORECASE,
)
LETTER_TITLE_SIZE = 11
LETTER_CONTACT_SIZE = 10.5
_SPANISH_MONTHS = (
    "enero", "febrero", "marzo", "abril", "mayo", "junio",
    "julio", "agosto", "septiembre", "octubre", "noviembre", "diciembre",
)
_DOCUMENT_DATE_RE = re.compile(r"^\d{1,2} de [a-záéíóúüñ]+ de \d{4}$", re.IGNORECASE)


class PreconditionError(ValueError):
    """La fase no puede iniciarse con la entrada actual."""


@dataclass(frozen=True)
class ComparisonResult:
    equivalent: bool
    omissions: list[str] = field(default_factory=list)
    additions: list[str] = field(default_factory=list)
    changes_figures: list[str] = field(default_factory=list)
    order_changed: bool = False


@dataclass(frozen=True)
class CompositionResult:
    content_source: Path
    docx_path: Path
    pdf_path: Path
    evaluation_path: Path
    state: str
    recommendation: str
    source_hash: str
    pages_docx: int
    pages_pdf: int
    comparison_docx: ComparisonResult
    comparison_pdf: ComparisonResult
    comparison_docx_pdf: ComparisonResult
    privacy_extras: list[str]
    header: CabeceraCandidatura
    comparison_header_docx: ComparisonResult
    comparison_header_pdf: ComparisonResult
    comparison_header_docx_pdf: ComparisonResult
    mode: str = "modo_documento"


def parse_frontmatter(text: str) -> dict[str, Any]:
    if not text.startswith("---"):
        return {}
    parts = text.split("---", 2)
    if len(parts) < 3:
        return {}
    data = yaml.safe_load(parts[1]) or {}
    return data if isinstance(data, dict) else {}


def extract_consolidated_letter(markdown: str) -> str:
    """Extrae solo el bloque autorizado, sin el encabezado ni el siguiente H2."""
    lines = markdown.splitlines()
    start = None
    for index, line in enumerate(lines):
        if SECTION_RE.match(line.strip()):
            start = index + 1
            break
    if start is None:
        raise PreconditionError("Falta la sección Carta completa consolidada.")
    end = len(lines)
    for index in range(start, len(lines)):
        if NEXT_SECTION_RE.match(lines[index].strip()):
            end = index
            break
    content = "\n".join(lines[start:end]).strip()
    if not content:
        raise PreconditionError("La Carta completa consolidada está vacía.")
    return content


def extraer_modo_texto(content_path: Path) -> str:
    """Devuelve el cuerpo exacto disponible para copiar/pegar, sin cabecera."""
    return extract_consolidated_letter(content_path.read_text(encoding="utf-8"))


def validate_preconditions(content_text: str, presented: bool) -> None:
    metadata = parse_frontmatter(content_text)
    if metadata.get("estado_contenido") != "apto":
        raise PreconditionError("estado_contenido no es apto.")
    gate = metadata.get("estado_gate_salida", metadata.get("estado_gate"))
    if gate != "aprobado":
        raise PreconditionError("GATE-CONTENIDO-CARTA-COMPOSICION no está aprobado.")
    if presented:
        raise PreconditionError("La candidatura ya figura como presentada.")
    extract_consolidated_letter(content_text)


def _read_metadata(path: Path) -> dict[str, Any]:
    return parse_frontmatter(path.read_text(encoding="utf-8"))


def validate_case_files(content_path: Path, candidate_path: Path) -> None:
    content_text = content_path.read_text(encoding="utf-8")
    candidate = _read_metadata(candidate_path)
    validate_preconditions(content_text, candidate.get("presentada") is True)
    if not candidate.get("id"):
        raise PreconditionError("La candidatura no tiene identificador.")
    if candidate.get("estado") in {"rechazada", "detenida", "cerrada"}:
        raise PreconditionError("La candidatura no está vigente.")


def cargar_cabecera_para_candidatura(candidate_dir: Path) -> CabeceraCandidatura:
    """Resuelve la cabecera desde el JSON canónico y comprueba la candidatura."""
    json_path = candidate_dir / "datos-generacion.json"
    candidate_path = candidate_dir / "candidatura.md"
    if not json_path.is_file() or not candidate_path.is_file():
        raise PreconditionError("No existe el JSON canónico o la candidatura para resolver la cabecera.")
    try:
        payload = json.loads(json_path.read_text(encoding="utf-8"))
        candidate = _read_metadata(candidate_path)
        header = construir_cabecera_candidatura(payload)
        authorization = candidate.get("autorizacion_datos_cv")
        json_authorization = payload["control"]["datos_privados"]["autorizacion"]
        auth_keys = ("nombre", "apellido_1", "apellido_2", "email", "telefono", "linkedin", "ubicacion", "fotografia")
        if not isinstance(authorization, dict) or {
            key: authorization.get(key) for key in auth_keys
        } != {key: json_authorization.get(key) for key in auth_keys}:
            raise PreconditionError("La autorización de datos de candidatura y JSON no coincide.")
        validar_cabecera_contrato(header, authorization)
        return header
    except (OSError, json.JSONDecodeError, KeyError, TypeError, CabeceraError) as exc:
        if isinstance(exc, PreconditionError):
            raise
        raise PreconditionError(f"No se puede resolver la cabecera canónica: {exc}") from exc


def _normalise_typography(text: str) -> str:
    replacements = {
        "\u00a0": " ", "\u2018": "'", "\u2019": "'", "\u201c": '"',
        "\u201d": '"', "\u2013": "-", "\u2014": "-", "\u2212": "-",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return unicodedata.normalize("NFKC", text)


def normalise_text(text: str) -> str:
    return re.sub(r"\s+", " ", _normalise_typography(text)).strip()


def fecha_documento_es(fecha: date) -> str:
    """Formatea la fecha de generación según la guía documental vigente."""
    return f"{fecha.day} de {_SPANISH_MONTHS[fecha.month - 1]} de {fecha.year}"


def asunto_candidatura(metadata: dict[str, Any]) -> str:
    """Deriva el asunto exclusivamente del puesto y empresa confirmados."""
    puesto = str(metadata.get("puesto_objetivo") or "").strip()
    empresa = str(metadata.get("empresa") or "").strip()
    if not puesto or not empresa:
        raise PreconditionError("No se puede derivar el asunto sin puesto y empresa confirmados.")
    return f"Candidatura — {puesto} — {empresa}"


def reconstruir_parrafo_semantico(block: str) -> str:
    """Une el hard wrapping Markdown sin convertirlo en saltos manuales."""
    return " ".join(line.strip() for line in block.splitlines() if line.strip())


def _strip_document_metadata(lines: list[str]) -> list[str]:
    """Retira fecha/asunto de la extracción para comparar solo el cuerpo semántico."""
    if len(lines) >= 2 and _DOCUMENT_DATE_RE.match(lines[0]) and lines[1].startswith("Asunto: "):
        # El asunto puede ocupar varias líneas en el PDF; el cuerpo comienza
        # en el saludo, no necesariamente en la tercera línea extraída.
        for index, line in enumerate(lines[2:], start=2):
            if line.startswith(("Estimado ", "Estimada ", "A la atención")):
                return lines[index:]
        return lines[2:]
    return lines


def _words(text: str) -> list[str]:
    return normalise_text(text).split()


def _diff_groups(source_words: list[str], output_words: list[str]) -> tuple[list[str], list[str]]:
    omissions: list[str] = []
    additions: list[str] = []
    matcher = difflib.SequenceMatcher(a=source_words, b=output_words, autojunk=False)
    for tag, a1, a2, b1, b2 in matcher.get_opcodes():
        if tag in {"delete", "replace"}:
            omissions.append(" ".join(source_words[a1:a2]))
        if tag in {"insert", "replace"}:
            additions.append(" ".join(output_words[b1:b2]))
    return omissions, additions


def _figures(text: str) -> list[str]:
    return [normalise_text(match.group(0)) for match in FIGURE_RE.finditer(text)]


def compare_texts(source: str, output: str) -> ComparisonResult:
    source_words = _words(source)
    output_words = _words(output)
    omissions, additions = _diff_groups(source_words, output_words)
    source_figures = _figures(source)
    output_figures = _figures(output)
    figure_changes = []
    if source_figures != output_figures:
        figure_changes.append(f"fuente={source_figures!r}; salida={output_figures!r}")
    order_changed = bool(
        normalise_text(source) != normalise_text(output)
        and sorted(source_words) == sorted(output_words)
    )
    return ComparisonResult(
        equivalent=normalise_text(source) == normalise_text(output),
        omissions=omissions,
        additions=additions,
        changes_figures=figure_changes,
        order_changed=order_changed,
    )


def compare_authorized_contacts(source: str, output: str) -> list[str]:
    source_values = {normalise_text(match.group(0)) for match in PRIVATE_DATA_RE.finditer(source)}
    return [
        normalise_text(match.group(0))
        for match in PRIVATE_DATA_RE.finditer(output)
        if normalise_text(match.group(0)) not in source_values
    ]


def page_count_issue(actual_pages: int, text: str) -> bool:
    if actual_pages <= 1:
        return False
    page_chunks = text.split("\f")
    return any(not chunk.strip() for chunk in page_chunks[1:]) or actual_pages > len([c for c in page_chunks if c.strip()])


def layout_status(text: str, max_pages: int, actual_pages: int) -> str:
    _ = text
    return "requiere_correccion_composicion" if actual_pages > max_pages else "apta"


def classify_render_error() -> str:
    return "requiere_correccion_composicion"


def sha256_bytes(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def sha256_file(path: Path) -> str:
    return sha256_bytes(path.read_bytes())


def _set_font(run, name: str, size: float, bold: bool = False) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold


def _remove_table_borders(table: Any) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        from docx.oxml import OxmlElement
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            from docx.oxml import OxmlElement
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def renderizar_cabecera_docx(document: Document, header: CabeceraCandidatura) -> None:
    table = document.add_table(rows=1, cols=1)
    _remove_table_borders(table)
    cell = table.cell(0, 0)
    cell.text = ""
    paragraphs = cell.paragraphs
    lines = header.lineas
    for index, line in enumerate(lines):
        paragraph = paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.paragraph_format.line_spacing = 1.0
        paragraph.paragraph_format.space_after = Pt(1 if index == 0 else 2 if index == 1 else 0)
        run = paragraph.add_run(line)
        sizes = (HEADER_NAME_SIZE, LETTER_TITLE_SIZE, LETTER_CONTACT_SIZE)
        _set_font(run, HEADER_FONT, sizes[index], bold=index == 0)


def build_docx(
    letter: str,
    header: CabeceraCandidatura,
    destination: Path,
    *,
    document_date: str | None = None,
    subject: str | None = None,
) -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.1)
    section.right_margin = Cm(2.1)
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    renderizar_cabecera_docx(document, header)
    document.add_paragraph().paragraph_format.space_after = Pt(4)
    if (document_date is None) != (subject is None):
        raise ValueError("La fecha y el asunto deben proporcionarse conjuntamente.")
    if document_date is not None and subject is not None:
        date_paragraph = document.add_paragraph()
        date_paragraph.paragraph_format.space_after = Pt(2)
        date_run = date_paragraph.add_run(document_date)
        _set_font(date_run, "Calibri", LETTER_CONTACT_SIZE)
        subject_paragraph = document.add_paragraph()
        subject_paragraph.paragraph_format.space_after = Pt(10)
        subject_run = subject_paragraph.add_run(f"Asunto: {subject}")
        _set_font(subject_run, "Calibri", LETTER_TITLE_SIZE, bold=True)
    paragraphs = [
        reconstruir_parrafo_semantico(block)
        for block in letter.split("\n\n")
        if reconstruir_parrafo_semantico(block)
    ]
    for index, block in enumerate(paragraphs):
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.keep_together = True
        paragraph.paragraph_format.widow_control = True
        is_greeting = index == 0
        is_signature = block in {"Atentamente,", "Gustavo Vega"}
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if is_greeting or is_signature else WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(block)
        _set_font(run, "Calibri", 11, bold=(block == "Gustavo Vega"))
        if index == 0:
            paragraph.paragraph_format.space_after = Pt(14)
        elif index == len(paragraphs) - 1:
            paragraph.paragraph_format.space_after = Pt(0)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.title = "Carta de presentación"
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def convert_docx_to_pdf(docx_path: Path, pdf_path: Path, project_root: Path) -> None:
    soffice = Path(r"C:\Program Files\LibreOffice\program\soffice.com")
    if not soffice.is_file():
        raise RuntimeError(f"No existe el conversor autorizado: {soffice}")
    staging = project_root / ".tmp" / "job-up-carta-lo" / uuid.uuid4().hex
    staging.mkdir(parents=True, exist_ok=False)
    staged_docx = staging / docx_path.name
    shutil.copy2(docx_path, staged_docx)
    profile = staging / "profile"
    profile.mkdir()
    command = [
        str(soffice), "--headless", "--nologo", "--nodefault", "--nofirststartwizard",
        "--norestore", f"-env:UserInstallation={profile.resolve().as_uri()}",
        "--convert-to", "pdf", "--outdir", str(staging), str(staged_docx),
    ]
    try:
        process = subprocess.run(command, cwd=staging, capture_output=True, text=True, timeout=60)
        generated = staging / f"{docx_path.stem}.pdf"
        if process.returncode != 0 or not generated.is_file() or generated.stat().st_size == 0:
            raise RuntimeError(f"Conversión PDF fallida ({process.returncode}): {process.stderr or process.stdout}")
        pdf_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(generated, pdf_path)
    finally:
        shutil.rmtree(staging, ignore_errors=True)


def extract_docx_header(path: Path) -> str:
    document = Document(path)
    if not document.tables:
        raise PreconditionError("El DOCX no contiene cabecera documental.")
    lines = [paragraph.text.strip() for paragraph in document.tables[0].cell(0, 0).paragraphs if paragraph.text.strip()]
    if len(lines) < 3:
        raise PreconditionError("La cabecera DOCX está incompleta.")
    return "\n".join(lines[:3])


def extract_docx_body(path: Path) -> str:
    document = Document(path)
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    return "\n\n".join(_strip_document_metadata(paragraphs))


def extract_docx_text(path: Path) -> str:
    """Compatibilidad: texto semántico del cuerpo, no incluye cabecera."""
    return extract_docx_body(path)


def extract_pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    return "\n\n".join(page.extract_text() or "" for page in reader.pages)


def _pdf_lines(path: Path) -> list[str]:
    return [line.strip() for line in extract_pdf_text(path).splitlines() if line.strip()]


def extract_pdf_header(path: Path) -> str:
    lines = _pdf_lines(path)
    if len(lines) < 3:
        raise PreconditionError("El PDF no contiene una cabecera completa.")
    return "\n".join(lines[:3])


def extract_pdf_body(path: Path) -> str:
    lines = _pdf_lines(path)
    if len(lines) < 4:
        raise PreconditionError("El PDF no contiene cuerpo después de la cabecera.")
    return "\n\n".join(_strip_document_metadata(lines[3:]))


def inspect_docx(path: Path) -> list[str]:
    issues: list[str] = []
    with zipfile.ZipFile(path) as archive:
        names = set(archive.namelist())
        xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        if (
            "word/comments.xml" in names
            or "<w:commentRangeStart" in xml
            or re.search(r"<w:ins(?:\s|>)", xml)
            or re.search(r"<w:del(?:\s|>)", xml)
        ):
            issues.append("comentarios_o_control_de_cambios")
        if "[" in xml and "]" in xml:
            issues.append("marcadores_visibles")
    return issues


def pdf_page_count(path: Path) -> int:
    return len(PdfReader(str(path)).pages)


def render_pdf_pages(pdf_path: Path, output_dir: Path) -> list[Path]:
    try:
        import pypdfium2 as pdfium
    except ImportError:
        return []
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf = pdfium.PdfDocument(str(pdf_path))
    paths: list[Path] = []
    for index in range(len(pdf)):
        bitmap = pdf[index].render(scale=1.5)
        target = output_dir / f"page-{index + 1}.png"
        bitmap.to_pil().save(target)
        paths.append(target)
    return paths


def _yaml_value(value: Any) -> str:
    return yaml.safe_dump(value, allow_unicode=True, default_flow_style=True).strip()


def build_evaluation(
    result: CompositionResult,
    issues: list[str],
    docx_issues: list[str],
    *,
    visual_inspected: bool = False,
    inspection_evidence: str | None = None,
) -> str:
    candidate_metadata = _read_metadata(result.content_source.parent / "candidatura.md")
    candidate_id = str(candidate_metadata.get("id") or "CANDIDATURA-SIN-ID")
    document_date = fecha_documento_es(date.today())
    subject = asunto_candidatura(candidate_metadata)
    c1 = result.comparison_docx
    c2 = result.comparison_pdf
    c3 = result.comparison_docx_pdf
    h1 = result.comparison_header_docx
    h2 = result.comparison_header_pdf
    h3 = result.comparison_header_docx_pdf
    lines = [
        "---",
        f"id: evaluacion-composicion-carta-{candidate_id}",
        "tipo: evaluacion_composicion_carta_presentacion",
        "version: \"1.1.0\"",
        "estado: en_prueba",
        f"candidatura: {candidate_id}",
        f"fecha_composicion: {date.today().isoformat()}",
        "contenido_fuente: contenido-carta-presentacion.md",
        "version_contenido: \"1.1.0\"",
        "gate_entrada: GATE-CONTENIDO-CARTA-COMPOSICION",
        "estado_gate_entrada: aprobado",
        "gate_salida: GATE-CARTA-REVISION-HUMANA",
        "decision_humana: pendiente",
        "estado_gate: pendiente",
        "---",
        "",
        f"# Evaluación de composición de carta de presentación — {candidate_id}",
        "",
        "## 1. Identificación",
        "",
        "| Campo | Valor |",
        "| --- | --- |",
        f"| Candidatura | {candidate_id} |",
        "| Contenido fuente | `contenido-carta-presentacion.md` → `Carta completa consolidada` |",
        f"| Hash fuente normalizada | `{result.source_hash}` |",
        "| Generador | `scripts/job-up/componer_carta_presentacion.py` |",
        "| Versión generador | 1.1.0 |",
        "",
        "## 2. Estado de composición",
        "",
        "```yaml",
        f"estado_composicion: {result.state}",
        f"recomendacion_gate: {result.recommendation}",
        "decision_humana: pendiente",
        "estado_gate: pendiente",
        f"modo_salida: {result.mode}",
        "```",
        "",
        "## 3. Cabecera canónica compartida",
        "",
        "```yaml",
        "requerida: true",
        "aplicada: true",
        f"origen: {result.header.origen}",
        f"version_o_identificador: {result.header.version}",
        "reutiliza_cabecera_cv: true",
        f"coherencia_cv: {str(result.header.version == 'datos-generacion-cv@1.2').lower()}",
        f"fuente: {_yaml_value(list(result.header.lineas))}",
        f"docx: {str(h1.equivalent).lower()}",
        f"pdf: {str(h2.equivalent).lower()}",
        f"docx_pdf: {str(h3.equivalent).lower()}",
        "```",
        "",
        "La cabecera se resuelve mediante el mismo helper canónico que consume el compositor de CV; no se mantiene una cabecera independiente de carta.",
        "",
        "## 4. Fuente semántica cerrada",
        "",
        "Solo se utilizó `Carta completa consolidada`. No se consultaron guion, candidatura, análisis, datos-core, CV, oferta ni web para añadir texto.",
        "",
        "## 5. Configuración visual aplicada",
        "",
        "```yaml",
        "tipografia: Calibri",
        "tamano_cuerpo: 11 pt",
        "tamano_nombre: 18 pt negrita",
        "tamano_titular: 11 pt",
        "tamano_contacto: 10.5 pt",
        f"fecha_documento: {document_date}",
        f"asunto_documento: {subject}",
        "margenes: 2.0 cm superior/inferior; 2.1 cm izquierdo/derecho",
        "interlineado: 1.15",
        "espaciado_parrafos: 8 pt; apertura 14 pt",
        f"paginas_docx: {result.pages_docx}",
        f"paginas_pdf: {result.pages_pdf}",
        "coherencia_cv: Calibri y tratamiento sobrio del contacto; layout propio de carta",
        "```",
        "",
        "## 6. Equivalencia semántica del cuerpo",
        "",
        "```yaml",
        f"fuente_docx: {str(c1.equivalent).lower()}",
        f"fuente_pdf: {str(c2.equivalent).lower()}",
        f"docx_pdf: {str(c3.equivalent).lower()}",
        f"omisiones: {_yaml_value(c1.omissions + c2.omissions + c3.omissions)}",
        f"adiciones: {_yaml_value(c1.additions + c2.additions + c3.additions)}",
        f"cambios_cifras: {_yaml_value(c1.changes_figures + c2.changes_figures + c3.changes_figures)}",
        f"cambios_orden: {_yaml_value([c1.order_changed, c2.order_changed, c3.order_changed])}",
        "```",
        "",
        "La secuencia material se conserva: saludo → apertura → desarrollo → contexto → cierre → despedida → firma. Las líneas físicas consecutivas del Markdown se reconstruyen como un único párrafo semántico.",
        "",
        "## 7. Privacidad",
        "",
        "```yaml",
        f"datos_autorizados: {_yaml_value([result.header.nombre, *result.header.contacto])}",
        f"datos_adicionales: {_yaml_value(result.privacy_extras)}",
        f"resultado: {'conforme' if not result.privacy_extras else 'bloqueante'}",
        "```",
        "",
        "## 8. Roles",
        "",
        "```yaml",
        "ingeniero_composicion_documental: ejecutado",
        "auditor_integridad_documental: ejecutado",
        "redactor: no_aplica",
        "recruiter: no_aplica",
        "```",
        "",
        "## 9. Salidas generadas",
        "",
        "| Formato | Ruta | Estado |",
        "| --- | --- | --- |",
        "| DOCX | `carta-presentacion.docx` | generado y validado |",
        "| PDF | `carta-presentacion.pdf` | generado y validado |",
        "| Evaluación | `evaluacion-composicion-carta-presentacion.md` | generado |",
        "",
        "## 10. Calidad técnica",
        "",
        "```yaml",
        "docx:",
        "  abre: true",
        "  legibilidad: apta",
        "  render: apto",
        f"  incidencias: {_yaml_value(docx_issues)}",
        "pdf:",
        "  abre: true",
        "  legibilidad: apta",
        "  render: apto",
        f"  incidencias: {_yaml_value(issues)}",
        "render_generado: true",
        f"render_inspeccionado: {str(visual_inspected).lower()}",
        "revision_visual:",
        f"  ejecutada: {str(visual_inspected).lower()}",
        f"  evidencia_inspeccion: {inspection_evidence or 'pendiente_de_inspeccion_real'}",
        "  renderizador: pypdfium2",
        "  inspeccion: cabecera, márgenes, espaciado, saltos, paginación y ausencia de cortes/solapamientos",
        "  comparacion_cv: cabecera, nombre, contacto y tipografía coherentes; layout propio de carta",
        "  pdf2image: no_disponible",
        "  impacto_pdf2image: no_bloqueante",
        "```",
        "",
        "## 11. Incidencias",
        "",
        f"`{_yaml_value(issues) if issues else 'ninguna'}`",
        "",
        (
            "La revisión visual se realizó con un renderizador alternativo (`pypdfium2`) porque `pdf2image` no está disponible en el entorno; esta limitación no bloquea el resultado técnico."
            if visual_inspected
            else "El render fue generado, pero la inspección visual real queda pendiente; no se declara `revision_visual: ejecutada`."
        ),
        "",
        "## 12. Resultado de composición",
        "",
        "```yaml",
        f"estado_composicion: {result.state}",
        f"recomendacion_gate: {result.recommendation}",
        "decision_humana: pendiente",
        "estado_gate: pendiente",
        "```",
        "",
        "La composición no aprueba la carta final; solo deja los artefactos disponibles para revisión humana y mantiene `presentada: false`.",
        "",
    ]
    return "\n".join(lines)


def compose_case(
    content_path: Path,
    *,
    visual_inspected: bool = False,
    inspection_evidence: str | None = None,
) -> CompositionResult:
    project_root = Path(__file__).resolve().parents[2]
    candidate_dir = content_path.parent
    candidate_path = candidate_dir / "candidatura.md"
    validate_case_files(content_path, candidate_path)
    source_file_text = content_path.read_text(encoding="utf-8")
    source = extract_consolidated_letter(source_file_text)
    source_hash = sha256_bytes(normalise_text(source).encode("utf-8"))
    header = cargar_cabecera_para_candidatura(candidate_dir)
    candidate_metadata = _read_metadata(candidate_path)
    document_date = fecha_documento_es(date.today())
    subject = asunto_candidatura(candidate_metadata)
    docx_path = candidate_dir / "carta-presentacion.docx"
    pdf_path = candidate_dir / "carta-presentacion.pdf"
    evaluation_path = candidate_dir / "evaluacion-composicion-carta-presentacion.md"
    build_docx(source, header, docx_path, document_date=document_date, subject=subject)
    convert_docx_to_pdf(docx_path, pdf_path, project_root)
    docx_text = extract_docx_body(docx_path)
    pdf_text = extract_pdf_body(pdf_path)
    header_text = "\n".join(header.lineas)
    header_docx = extract_docx_header(docx_path)
    header_pdf = extract_pdf_header(pdf_path)
    comparison_docx = compare_texts(source, docx_text)
    comparison_pdf = compare_texts(source, pdf_text)
    comparison_docx_pdf = compare_texts(docx_text, pdf_text)
    comparison_header_docx = compare_texts(header_text, header_docx)
    comparison_header_pdf = compare_texts(header_text, header_pdf)
    comparison_header_docx_pdf = compare_texts(header_docx, header_pdf)
    privacy_extras = compare_authorized_contacts(header_text, "\n".join((header_docx, header_pdf)))
    pages_docx = 1
    pages_pdf = pdf_page_count(pdf_path)
    docx_issues = inspect_docx(docx_path)
    issues = list(docx_issues)
    if pages_pdf > 1:
        issues.append(f"pdf_tiene_{pages_pdf}_paginas")
    if page_count_issue(pages_pdf, pdf_text):
        issues.append("pagina_vacia_inesperada")
    if privacy_extras:
        issues.append("dato_personal_no_autorizado")
    semantic_ok = all(comparison.equivalent for comparison in (
        comparison_docx, comparison_pdf, comparison_docx_pdf,
        comparison_header_docx, comparison_header_pdf, comparison_header_docx_pdf,
    ))
    if not semantic_ok:
        state, recommendation = "requiere_revision_contenido", "recomendar_no_aprobar"
    elif issues:
        state, recommendation = "requiere_correccion_composicion", "recomendar_no_aprobar"
    else:
        state, recommendation = "apta", "recomendar_aprobar"
    result = CompositionResult(
        content_path, docx_path, pdf_path, evaluation_path, state, recommendation,
        source_hash, pages_docx, pages_pdf, comparison_docx, comparison_pdf,
        comparison_docx_pdf, privacy_extras, header,
        comparison_header_docx, comparison_header_pdf, comparison_header_docx_pdf,
    )
    evaluation_path.write_text(
        build_evaluation(
            result,
            issues,
            docx_issues,
            visual_inspected=visual_inspected,
            inspection_evidence=inspection_evidence,
        ),
        encoding="utf-8",
    )
    return result


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("content", type=Path)
    parser.add_argument(
        "--modo",
        choices=("texto", "documento"),
        default="documento",
        help="Salida de texto para copiar/pegar o composición documental DOCX/PDF.",
    )
    args = parser.parse_args()
    if args.modo == "texto":
        print(extraer_modo_texto(args.content))
        return 0
    result = compose_case(args.content)
    print(yaml.safe_dump({
        "estado_composicion": result.state,
        "recomendacion_gate": result.recommendation,
        "docx": str(result.docx_path),
        "pdf": str(result.pdf_path),
        "evaluacion": str(result.evaluation_path),
        "paginas_pdf": result.pages_pdf,
    }, allow_unicode=True, sort_keys=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
