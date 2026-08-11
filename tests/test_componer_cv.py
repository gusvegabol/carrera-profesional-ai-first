import copy
import importlib.util
import json
import sys
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
MODULE_PATH = ROOT / "scripts/job-up/componer_cv.py"
sys.path.insert(0, str(ROOT / "scripts/job-up"))
FIXTURE = ROOT / "boveda-entrevista-profesional/busqueda-empleo/candidaturas/CAND-2026-020-lidl-responsable-turno-tienda-tamaraceite/datos-generacion.json"
TEMPLATE = ROOT / "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas/TEMPLATE_CV_FORMATO.docx"
PHOTO = ROOT / "boveda-entrevista-profesional/busqueda-empleo/foto-perfil.png"


def load_module():
    spec = importlib.util.spec_from_file_location("componer_cv", MODULE_PATH)
    module = importlib.util.module_from_spec(spec)
    sys.modules[spec.name] = module
    spec.loader.exec_module(module)
    return module


class CompositorCVTests(unittest.TestCase):
    def setUp(self):
        self.module = load_module()
        self.payload = json.loads(FIXTURE.read_text(encoding="utf-8"))

    def test_modelo_respeta_orden_y_textos_literales(self):
        data = copy.deepcopy(self.payload)
        data["contenido_cv"]["encabezado"]["unidades"].append({
            "id_contenido": "C-099", "tipo": "subtitular", "orden": 0,
            "texto": "Primero & literal", "trazabilidad": {},
        })
        data["contenido_cv"]["secciones"] = list(reversed(data["contenido_cv"]["secciones"]))

        model = self.module.construir_modelo_cv(data)

        self.assertEqual(model.encabezado.unidades[0].texto, "Primero & literal")
        self.assertEqual([section.id_seccion for section in model.secciones], ["SEC-01", "SEC-02", "SEC-03", "SEC-04"])
        self.assertEqual(model.secciones[1].bloques[0].unidades[0].texto, self.payload["contenido_cv"]["secciones"][1]["bloques"][0]["unidades"][0]["texto"])

    def test_modelo_ignora_metadatos_y_control(self):
        expected = self.module.construir_modelo_cv(self.payload)
        changed = copy.deepcopy(self.payload)
        changed["candidatura"]["empresa"] = "Otra empresa"
        changed["control"] = {"contenido_ajeno": "No debe aparecer"}
        changed["generacion"]["guion_origen"] = "otra-ruta.md"

        self.assertEqual(self.module.construir_modelo_cv(changed), expected)

    def test_latex_es_dinamico_y_preserva_texto_visible(self):
        model = self.module.construir_modelo_cv(self.payload)

        latex = self.module.renderizar_latex(model)

        self.assertIn(r"Gustavo Vega", latex)
        self.assertIn(r"Reduje un 30 \% las caducidades", latex)
        self.assertLess(latex.index("Perfil"), latex.index("Experiencia operativa relevante"))
        self.assertNotIn("[EXPERIENCIA", latex)
        self.assertEqual(latex.count("{"), latex.count("}"))

    def test_docx_es_dinamico_e_incluye_fotografia(self):
        model = self.module.construir_modelo_cv(self.payload)
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "cv.docx"

            self.module.renderizar_docx(model, TEMPLATE, output, PHOTO)

            document = Document(output)
            text = "\n".join(p.text for p in document.paragraphs)
            table_text = "\n".join(cell.text for table in document.tables for row in table.rows for cell in row.cells)
            self.assertIn("Perfil", text)
            self.assertIn("Reduje un 30 % las caducidades", text)
            self.assertIn("Gustavo Vega", table_text)
            self.assertEqual(len(document.inline_shapes), 1)
            self.assertNotIn("[EXPERIENCIA", text + table_text)

    def test_cabecera_cv_procede_del_helper_canonico_compartido(self):
        header = self.module.construir_cabecera_candidatura(self.payload)
        model = self.module.construir_modelo_cv(self.payload)
        self.assertEqual(model.encabezado.nombre, header.nombre)
        self.assertEqual(model.encabezado.unidades[0].texto, header.titular)
        self.assertEqual(tuple(item.texto for item in model.encabezado.contacto), header.contacto)


if __name__ == "__main__":
    unittest.main()
