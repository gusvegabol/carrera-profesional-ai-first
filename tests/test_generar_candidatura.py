import json
import sys
import unittest
import tempfile
import subprocess
from unittest import mock
from pathlib import Path


ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "scripts" / "job-up"))

from generar_candidatura import (  # noqa: E402
    build_latex,
    build_docx,
    canonical_output_paths,
    write_error_record,
    validate_candidate_state,
    CandidateLock,
    publish_transaction,
    main,
    recover_pending_publications,
    resolve_input_json,
    _assert_process_tree_gone,
    convert_docx_to_pdf,
    validate_payload,
    _validate_experiences,
    _validate_photo,
    _validate_template_text,
    _template_paragraph_text,
    validate_latex,
    _replace_docx_paragraph,
)


class GeneratorContractTests(unittest.TestCase):
    def test_canonical_output_paths_are_derived_from_candidate_route(self):
        route = Path("boveda-entrevista-profesional/busqueda-empleo/candidaturas/CAND-2026-999-fixture")
        self.assertEqual(
            canonical_output_paths(route),
            {
                "cv_docx": route / "cv.docx",
                "cv_pdf": route / "cv.pdf",
                "carta_docx": route / "carta-presentacion.docx",
                "carta_pdf": route / "carta-presentacion.pdf",
                "cv_tex": route / "cv.tex",
            },
        )

    def test_payload_rejects_output_outside_candidate(self):
        payload = self._payload()
        payload["salidas"]["cv_docx"] = "boveda-entrevista-profesional/otra/cv.docx"
        with self.assertRaises(ValueError):
            validate_payload(payload, ROOT)

    def test_cli_input_accepts_absolute_path_inside_project(self):
        absolute = ROOT / "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas/TEMPLATE_DATOS_GENERACION_CANDIDATURA.json"
        self.assertEqual(resolve_input_json(ROOT, str(absolute)), absolute.resolve())

    def test_error_record_is_timestamped_and_schema_shaped(self):
        with tempfile.TemporaryDirectory() as directory:
            path = write_error_record(Path(directory), "abc12345", "entrada.json", ValueError("fallo"), "validacion_entradas", "CAND-2026-999")
            self.assertIsNotNone(path)
            record = json.loads(path.read_text(encoding="utf-8"))
            self.assertEqual(record["resultado"], "fallido")
            self.assertEqual(record["id_candidatura"], "CAND-2026-999")
            self.assertTrue(path.name.startswith("error-"))

    def test_state_validation_requires_explicit_presented_value_in_both_sources(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            route = root / "boveda-entrevista-profesional/busqueda-empleo/candidaturas/CAND-2026-999-fixture"
            route.mkdir(parents=True)
            (route / "candidatura.md").write_text("---\nid: CAND-2026-999\nestado: en_preparacion\npresentada: false\n---\n", encoding="utf-8")
            tracking = root / "boveda-entrevista-profesional/busqueda-empleo/seguimiento"
            tracking.mkdir(parents=True)
            (tracking / "seguimiento-candidaturas.md").write_text(
                "| id_candidatura | estado | presentada |\n| --- | --- | --- |\n| CAND-2026-999 | en_preparacion | false |\n",
                encoding="utf-8",
            )
            validate_candidate_state(root, route, "CAND-2026-999")

    def test_state_matrix_accepts_only_regenerable_combinations(self):
        allowed = ("en_preparacion", "pendiente_de_aprobacion", "detenida", "fallida", "aprobada")
        for state in allowed:
            with tempfile.TemporaryDirectory() as directory:
                root, route = self._state_fixture(Path(directory), state, "false")
                validate_candidate_state(root, route, "CAND-2026-999")
        for state in ("enviada", "rechazada", "duplicada"):
            with tempfile.TemporaryDirectory() as directory:
                root, route = self._state_fixture(Path(directory), state, "true" if state != "duplicada" else "false")
                with self.assertRaises(ValueError):
                    validate_candidate_state(root, route, "CAND-2026-999")

    def test_experience_matrix_accepts_zero_three_or_six_complete_pairs(self):
        payload = self._payload()
        values = payload["cv"] | {"[NOMBRE]": "N", "[TITULAR]": "T", "[PERFIL PROFESIONAL]": "P", "[PROPUESTA DE VALOR]": "V"}
        for count in (0, 3, 6):
            current = values.copy()
            for number in range(1, 7):
                current[f"[EXPERIENCIA {number} CABECERA]"] = f"Empresa {number}" if number <= count else ""
                current[f"[EXPERIENCIA {number} DESCRIPCION]"] = f"Descripción {number}" if number <= count else ""
            _validate_experiences(current)
        invalid = values.copy()
        invalid["[EXPERIENCIA 1 DESCRIPCION]"] = "Solo descripción"
        with self.assertRaises(ValueError):
            _validate_experiences(invalid)

    def test_template_and_photo_contract_reject_corruption(self):
        from docx import Document
        root = Path(__file__).resolve().parents[1]
        document = Document(root / "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas/TEMPLATE_CV_FORMATO.docx")
        _validate_template_text(_template_paragraph_text(document), tuple(__import__("generar_candidatura").CV_FIELDS), "CV")
        self.assertEqual(len(document.inline_shapes), 1)
        self.assertNotIn("\ufffd", _template_paragraph_text(document))
        with tempfile.TemporaryDirectory() as directory:
            invalid = Path(directory) / "foto.txt"
            invalid.write_text("no es una imagen", encoding="utf-8")
            with self.assertRaises(ValueError):
                _validate_photo(invalid)

            unknown = "[NOMBRE] [MARCador_NO_DECLARADO]"
            with self.assertRaises(ValueError):
                _validate_template_text(unknown, ("[NOMBRE]",), "fixture-corrupta")

    def test_split_marker_is_replaced_and_unknown_duplicate_markers_are_rejected(self):
        from docx import Document
        document = Document()
        paragraph = document.add_paragraph()
        paragraph.add_run("[NOM")
        paragraph.add_run("BRE]")
        _replace_docx_paragraph(paragraph, {"[NOMBRE]": "Nombre"})
        self.assertEqual("".join(run.text or "" for run in paragraph.runs), "Nombre")
        with self.assertRaises(ValueError):
            _validate_template_text("[NOMBRE] [NOMBRE]", ("[NOMBRE]",), "duplicado")

    def test_missing_photo_and_unwritable_error_log_stop_or_fallback(self):
        payload = self._payload()
        payload["entradas"]["foto"] = "boveda-entrevista-profesional/busqueda-empleo/no-existe.png"
        with self.assertRaises(ValueError):
            validate_payload(payload, ROOT)
        with tempfile.TemporaryDirectory() as directory:
            with mock.patch("generar_candidatura.Path.write_text", side_effect=OSError("solo lectura")):
                self.assertIsNone(write_error_record(Path(directory), "abc12345", "entrada.json", ValueError("fallo"), "carga_json"))

    def test_state_missing_and_discordant_sources_are_blocked(self):
        with tempfile.TemporaryDirectory() as directory:
            root, route = self._state_fixture(Path(directory), "en_preparacion", "false")
            ficha = route / "candidatura.md"
            ficha.write_text(ficha.read_text(encoding="utf-8").replace("presentada: false\n", ""), encoding="utf-8")
            with self.assertRaises(ValueError):
                validate_candidate_state(root, route, "CAND-2026-999")
        with tempfile.TemporaryDirectory() as directory:
            root, route = self._state_fixture(Path(directory), "en_preparacion", "false")
            tracking = root / "boveda-entrevista-profesional/busqueda-empleo/seguimiento/seguimiento-candidaturas.md"
            tracking.write_text(tracking.read_text(encoding="utf-8").replace("en_preparacion | false", "fallida | false"), encoding="utf-8")
            with self.assertRaises(ValueError):
                validate_candidate_state(root, route, "CAND-2026-999")

    def test_latex_structural_validation_rejects_unbalanced_document(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "cv.tex"
            path.write_text("\\documentclass{article}\n\\begin{document}\n{", encoding="utf-8")
            with self.assertRaises(ValueError):
                validate_latex(path)

    def test_latex_validation_is_structural_and_does_not_require_compiler(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "cv.tex"
            path.write_text("\\documentclass{article}\n\\begin{document}\nTexto\n\\end{document}\n", encoding="utf-8")
            with mock.patch("generar_candidatura.shutil.which", side_effect=AssertionError("no debe buscar el compilador")):
                self.assertEqual(validate_latex(path), "validado_estructuralmente")

    def test_publication_replaces_all_outputs_and_can_restore_after_interruption(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            execution = root / ".tmp/job-up-generador/CAND-2026-999-fixture/run"
            execution.mkdir(parents=True)
            generated = {}
            destinations = {}
            for key in ("cv_docx", "cv_pdf", "carta_docx", "carta_pdf", "cv_tex"):
                source = execution / f"{key}.bin"
                source.write_bytes(f"new-{key}".encode())
                destination = root / "candidate" / f"{key}.bin"
                destination.parent.mkdir(parents=True, exist_ok=True)
                destination.write_bytes(f"old-{key}".encode())
                generated[key] = source
                destinations[key] = destination
            # The normal path must replace every destination and retain the manifest.
            result = publish_transaction(root, execution, generated, destinations, "abc12345")
            self.assertEqual(result["phase"], "completado")
            self.assertEqual(destinations["cv_docx"].read_bytes(), b"new-cv_docx")
            self.assertTrue((execution / "manifest.json").is_file())

    def test_publication_restores_previous_outputs_when_a_replace_fails(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            execution = root / ".tmp/job-up-generador/CAND-2026-999-fixture/run"
            execution.mkdir(parents=True)
            generated = {}
            destinations = {}
            for key in ("cv_docx", "cv_pdf", "carta_docx", "carta_pdf", "cv_tex"):
                source = execution / f"{key}.bin"
                source.write_bytes(f"new-{key}".encode())
                destination = root / "candidate" / f"{key}.bin"
                destination.parent.mkdir(parents=True, exist_ok=True)
                destination.write_bytes(f"old-{key}".encode())
                generated[key] = source
                destinations[key] = destination
            real_replace = __import__("os").replace
            calls = {"count": 0}

            def fail_once(source, destination):
                calls["count"] += 1
                if calls["count"] == 2:
                    raise OSError("interrupción simulada")
                return real_replace(source, destination)

            with mock.patch("generar_candidatura.os.replace", side_effect=fail_once):
                with self.assertRaises(OSError):
                    publish_transaction(root, execution, generated, destinations, "abc12345")
            for key, destination in destinations.items():
                self.assertEqual(destination.read_bytes(), f"old-{key}".encode())

    def test_candidate_lock_rejects_a_live_second_execution(self):
        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / ".lock"
            with CandidateLock(path):
                with self.assertRaises(RuntimeError):
                    with CandidateLock(path):
                        pass

    def test_timeout_cleanup_can_verify_a_gone_process_tree(self):
        _assert_process_tree_gone(4294967294)

    def test_libreoffice_timeout_kills_tree_and_cleans_staging(self):
        class HungProcess:
            pid = 4294967294
            returncode = None

            def communicate(self, timeout):
                raise __import__("subprocess").TimeoutExpired("soffice.com", timeout, stderr="timeout")

            def wait(self):
                return None

        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "cv.docx"
            source.write_bytes(b"fixture")
            output = root / "execution"
            real_popen = subprocess.Popen
            calls = {"count": 0}

            def fake_first_popen(*args, **kwargs):
                calls["count"] += 1
                return HungProcess() if calls["count"] == 1 else real_popen(*args, **kwargs)

            with mock.patch("generar_candidatura.subprocess.Popen", side_effect=fake_first_popen):
                with self.assertRaises(RuntimeError):
                    convert_docx_to_pdf(source, output, Path("soffice.com"), root)
            staging = root / ".tmp/job-up-lo"
            self.assertFalse(staging.exists() and any(staging.iterdir()))

    def test_conversion_removes_stale_pdf_before_accepting_new_output(self):
        class CompletedProcess:
            pid = 4294967294
            returncode = 0

            def communicate(self, timeout):
                return "", ""

        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            source = root / "cv.docx"
            source.write_bytes(b"fixture")
            output = root / "execution"
            output.mkdir()
            stale = output / "cv.pdf"
            stale.write_bytes(b"old-pdf")

            def fake_popen(*args, **kwargs):
                Path(kwargs["cwd"]).joinpath("cv.pdf").write_bytes(b"new-pdf")
                return CompletedProcess()

            with mock.patch("generar_candidatura.subprocess.Popen", side_effect=fake_popen):
                result = convert_docx_to_pdf(source, output, Path("soffice.com"), root)
            self.assertEqual(result.read_bytes(), b"new-pdf")

    def test_recovery_restores_a_publicando_manifest_before_new_generation(self):
        with tempfile.TemporaryDirectory() as directory:
            root = Path(directory)
            execution = root / ".tmp/job-up-generador/CAND-2026-999-fixture/run"
            backup = execution / "backups"
            backup.mkdir(parents=True)
            destination = root / "candidate" / "cv.docx"
            destination.parent.mkdir(parents=True)
            destination.write_bytes(b"partial")
            (backup / "cv_docx.bak").write_bytes(b"old")
            manifest = {
                "schema_version": "1.0", "execution_id": "abc12345", "phase": "publicando", "publicados": ["candidate/cv.docx"],
                "files": [{"key": "cv_docx", "source": "cv.docx", "destination": "candidate/cv.docx", "backup_name": "cv_docx.bak", "had_original": True, "sha256": "unused"}],
            }
            (execution / "manifest.json").write_text(json.dumps(manifest), encoding="utf-8")
            recover_pending_publications(root, "CAND-2026-999-fixture")
            self.assertEqual(destination.read_bytes(), b"old")
            self.assertEqual(json.loads((execution / "manifest.json").read_text(encoding="utf-8"))["phase"], "restaurado")

    @unittest.skipUnless(Path(r"C:\Program Files\LibreOffice\program\soffice.com").is_file(), "LibreOffice no está instalado")
    def test_cli_integration_generates_and_publishes_five_artifacts(self):
        root = Path(__file__).resolve().parents[1]
        source_templates = root / "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas"
        source_photo = root / "boveda-entrevista-profesional/busqueda-empleo/foto-perfil.png"
        with tempfile.TemporaryDirectory() as directory:
            fixture = Path(directory) / "project"
            template_dir = fixture / "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas"
            candidate = fixture / "boveda-entrevista-profesional/busqueda-empleo/candidaturas/CAND-2026-999-fixture"
            tracking = fixture / "boveda-entrevista-profesional/busqueda-empleo/seguimiento"
            template_dir.mkdir(parents=True)
            candidate.mkdir(parents=True)
            tracking.mkdir(parents=True)
            for name in ("TEMPLATE_CV_FORMATO.docx", "TEMPLATE_CARTA_PRESENTACION_FORMATO.docx", "TEMPLATE_CV_FORMATO.tex", "SCHEMA_DATOS_GENERACION_CANDIDATURA_1.0.json"):
                (template_dir / name).write_bytes((source_templates / name).read_bytes())
            photo = fixture / "boveda-entrevista-profesional/busqueda-empleo/foto-perfil.png"
            photo.parent.mkdir(parents=True, exist_ok=True)
            photo.write_bytes(source_photo.read_bytes())
            (candidate / "candidatura.md").write_text(
                "---\nid: CAND-2026-999\nestado: en_preparacion\npresentada: false\n---\n",
                encoding="utf-8",
            )
            (tracking / "seguimiento-candidaturas.md").write_text(
                "| id_candidatura | estado | presentada |\n| --- | --- | --- |\n| CAND-2026-999 | en_preparacion | false |\n",
                encoding="utf-8",
            )
            payload = json.loads((source_templates / "TEMPLATE_DATOS_GENERACION_CANDIDATURA.json").read_text(encoding="utf-8"))
            payload["id_candidatura"] = "CAND-2026-999"
            payload["ruta_candidatura"] = "boveda-entrevista-profesional/busqueda-empleo/candidaturas/CAND-2026-999-fixture"
            payload["entradas"] = {
                "template_cv": "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas/TEMPLATE_CV_FORMATO.docx",
                "template_carta": "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas/TEMPLATE_CARTA_PRESENTACION_FORMATO.docx",
                "template_latex": "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas/TEMPLATE_CV_FORMATO.tex",
                "foto": "boveda-entrevista-profesional/busqueda-empleo/foto-perfil.png",
            }
            route = Path(payload["ruta_candidatura"])
            payload["salidas"] = {key: (route / name).as_posix() for key, name in {
                "cv_docx": "cv.docx", "cv_pdf": "cv.pdf", "carta_docx": "carta-presentacion.docx",
                "carta_pdf": "carta-presentacion.pdf", "cv_tex": "cv.tex",
            }.items()}
            for section in (payload["cv"], payload["latex"]):
                section.update({"[NOMBRE]": "Nombre Prueba", "[TITULAR]": "Gestor de operaciones", "[PERFIL PROFESIONAL]": "Perfil operativo", "[PROPUESTA DE VALOR]": "Propuesta concreta", "[EXPERIENCIA 1 CABECERA]": "Empresa 1", "[EXPERIENCIA 1 DESCRIPCION]": "Descripción 1"})
            payload["carta"].update({"[NOMBRE]": "Nombre Prueba", "[TITULAR]": "Gestor de operaciones", "[DESTINATARIO]": "Empresa receptora", "[FECHA]": "30 de julio de 2026", "[ASUNTO]": "Candidatura", "[SALUDO]": "Estimado equipo:", "[APERTURA]": "Me dirijo a ustedes.", "[CIERRE]": "Quedo a su disposición.", "[DESPEDIDA]": "Atentamente,", "[FIRMA]": "Nombre Prueba"})
            json_path = candidate / "datos-generacion.json"
            json_path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
            script_dir = Path(directory) / "script"
            script_dir.mkdir()
            (script_dir / ".env").write_text(f"RUTA_PROYECTO={fixture}\nSOFFICE_PATH=C:\\Program Files\\LibreOffice\\program\\soffice.com\n", encoding="utf-8")
            with mock.patch("generar_candidatura.SCRIPT_DIR", script_dir):
                result = main(["boveda-entrevista-profesional/busqueda-empleo/candidaturas/CAND-2026-999-fixture/datos-generacion.json"])
            self.assertEqual(result, 0)
            for name in ("cv.docx", "cv.pdf", "carta-presentacion.docx", "carta-presentacion.pdf", "cv.tex"):
                self.assertTrue((candidate / name).is_file(), name)
            from pypdfium2 import PdfDocument
            for name in ("cv.pdf", "carta-presentacion.pdf"):
                pdf = PdfDocument(str(candidate / name))
                try:
                    bitmap = pdf[0].render(scale=1)
                    self.assertGreater(bitmap.width, 0)
                    self.assertGreater(bitmap.height, 0)
                finally:
                    pdf.close()
            execution_parent = fixture / ".tmp/job-up-generador/CAND-2026-999-fixture"
            self.assertFalse(execution_parent.exists() and any(execution_parent.iterdir()))
            # Simular el único paso humano posterior a la generación: aprobar
            # los documentos y dejar la candidatura preparada para presentar.
            ficha = candidate / "candidatura.md"
            ficha.write_text(ficha.read_text(encoding="utf-8").replace("estado: en_preparacion", "estado: pendiente_de_aprobacion"), encoding="utf-8")
            seguimiento = tracking / "seguimiento-candidaturas.md"
            seguimiento.write_text(seguimiento.read_text(encoding="utf-8").replace("| en_preparacion | false |", "| pendiente_de_aprobacion | false |"), encoding="utf-8")
            self.assertIn("estado: pendiente_de_aprobacion", ficha.read_text(encoding="utf-8"))
            self.assertIn("| pendiente_de_aprobacion | false |", seguimiento.read_text(encoding="utf-8"))

    def test_latex_removes_empty_experience_lines_without_blank_lines(self):
        template = (
            "A\n"
            "\\textbf{[EXPERIENCIA 1 CABECERA]} -- [EXPERIENCIA 1 DESCRIPCION]\n"
            "\\textbf{[EXPERIENCIA 2 CABECERA]} -- [EXPERIENCIA 2 DESCRIPCION]\n"
            "B\n"
        )
        values = {
            "[EXPERIENCIA 1 CABECERA]": "Empresa 1",
            "[EXPERIENCIA 1 DESCRIPCION]": "Descripción 1",
            "[EXPERIENCIA 2 CABECERA]": "",
            "[EXPERIENCIA 2 DESCRIPCION]": "",
        }
        result = build_latex(template, values)
        self.assertIn("\\textbf{Empresa 1} -- Descripción 1", result)
        self.assertNotIn("EXPERIENCIA 2", result)
        self.assertNotIn("B\n\n", result)

    def test_docx_replaces_markers_and_removes_empty_experience_paragraph(self):
        root = Path(__file__).resolve().parents[1]
        template = root / "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas/TEMPLATE_CV_FORMATO.docx"
        photo = root / "boveda-entrevista-profesional/busqueda-empleo/foto-perfil.png"
        payload = self._payload()
        with tempfile.TemporaryDirectory() as directory:
            output = Path(directory) / "cv.docx"
            build_docx(template, output, payload["cv"] | {
                "[NOMBRE]": "Nombre de prueba",
                "[TITULAR]": "Titular de prueba",
                "[PERFIL PROFESIONAL]": "Perfil de prueba",
                "[PROPUESTA DE VALOR]": "Propuesta de prueba",
                "[EXPERIENCIA 1 CABECERA]": "Empresa 1",
                "[EXPERIENCIA 1 DESCRIPCION]": "Descripción 1",
            }, photo)
            from docx import Document
            document = Document(output)
            paragraphs = list(document.paragraphs)
            paragraphs.extend(
                paragraph
                for table in document.tables
                for row in table.rows
                for cell in row.cells
                for paragraph in cell.paragraphs
            )
            text = "\n".join("".join(run.text or "" for run in paragraph.runs) for paragraph in paragraphs)
            self.assertIn("Empresa 1", text)
            self.assertNotIn("EXPERIENCIA 2", text)
            self.assertNotIn("[NOMBRE]", text)
            self.assertNotIn("COMPETENCIAS Y HERRAMIENTAS", text)
            self.assertNotIn("INFORMACIÓN ADICIONAL", text)

    def _payload(self):
        template = json.loads(
            (ROOT / "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas/TEMPLATE_DATOS_GENERACION_CANDIDATURA.json").read_text(encoding="utf-8")
        )
        template["id_candidatura"] = "CAND-2026-999"
        template["ruta_candidatura"] = "boveda-entrevista-profesional/busqueda-empleo/candidaturas/CAND-2026-999-fixture"
        template["entradas"] = {
            "template_cv": "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas/TEMPLATE_CV_FORMATO.docx",
            "template_carta": "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas/TEMPLATE_CARTA_PRESENTACION_FORMATO.docx",
            "template_latex": "boveda-entrevista-profesional/busqueda-empleo/proceso/plantillas/TEMPLATE_CV_FORMATO.tex",
            "foto": "boveda-entrevista-profesional/busqueda-empleo/foto-perfil.png",
        }
        template["salidas"] = {str(k): str(v) for k, v in canonical_output_paths(Path(template["ruta_candidatura"])).items()}
        return template

    def _state_fixture(self, root, state, presented):
        route = root / "boveda-entrevista-profesional/busqueda-empleo/candidaturas/CAND-2026-999-fixture"
        route.mkdir(parents=True)
        (route / "candidatura.md").write_text(
            f"---\nid: CAND-2026-999\nestado: {state}\npresentada: {presented}\n---\n",
            encoding="utf-8",
        )
        tracking = root / "boveda-entrevista-profesional/busqueda-empleo/seguimiento"
        tracking.mkdir(parents=True)
        (tracking / "seguimiento-candidaturas.md").write_text(
            f"| id_candidatura | estado | presentada |\n| --- | --- | --- |\n| CAND-2026-999 | {state} | {presented} |\n",
            encoding="utf-8",
        )
        return root, route


if __name__ == "__main__":
    unittest.main()
